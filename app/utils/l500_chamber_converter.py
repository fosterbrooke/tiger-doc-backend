from docx import Document
from docx.shared import RGBColor, Pt
from lxml import etree
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import zipfile
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH, WD_UNDERLINE
import sys
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import re
import os
from docx.oxml.shared import OxmlElement, qn as oxml_qn

def delete_table_row(document: Document, table_index, row_index: int) -> None:
    table = document.tables[table_index]
    try:
        table_element = table._tbl
        rows = table_element.xpath('.//w:tr')
        
        if not rows:
            raise ValueError("Table has no rows to delete")
            
        if row_index < 0 or row_index >= len(rows):
            raise IndexError(f"Row index {row_index} is out of range. Table has {len(rows)} rows.")
        
        row_to_delete = rows[row_index]
        table_element.remove(row_to_delete)
        
    except Exception as e:
        raise RuntimeError(f"Failed to delete row: {str(e)}") from e

def extract_specific_textbox_text(docx_file, textbox_index):
    try:
        with zipfile.ZipFile(docx_file) as z:
            xml_content = z.read('word/document.xml')
        tree = etree.fromstring(xml_content)
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
              'v': 'urn:schemas-microsoft-com:vml'}

        textboxes = tree.xpath('//v:shape[v:textbox]', namespaces=ns)

        if textbox_index < 0 or textbox_index >= len(textboxes):
            return None

        target_textbox = textboxes[textbox_index]
        text_elements = target_textbox.xpath('.//w:t', namespaces=ns)
        text = "".join([t.text for t in text_elements])

        return text

    except Exception as e:
        print(f"Error: {e}")
        return None

def copy_cell_content_to_target_cell(source_doc, source_table_index, source_row_index, source_col_index,
                                    target_doc, target_table_index, target_row_index, target_col_index):
    """
    Copies cell content with formatting from one cell to another cell in a target Document object.
    """
    try:
        # ===== TARGET CELL VALIDATION =====
        if not hasattr(target_doc, 'tables') or target_table_index >= len(target_doc.tables) or target_table_index < 0:
            print(f"Error: Target table index {target_table_index} out of range (0-{len(target_doc.tables)-1})")
            return False
            
        target_table = target_doc.tables[target_table_index]
        if target_row_index >= len(target_table.rows) or target_row_index < 0:
            print(f"Error: Target row index {target_row_index} out of range (0-{len(target_table.rows)-1})")
            return False
            
        target_row = target_table.rows[target_row_index]
        if target_col_index >= len(target_row.cells) or target_col_index < 0:
            print(f"Error: Target column index {target_col_index} out of range (0-{len(target_row.cells)-1})")
            return False
            
        target_cell = target_row.cells[target_col_index]
        
        # ===== SOURCE CONTENT EXTRACTION =====
        try:
            source_xml_content = source_doc.part.blob
            source_tree = etree.fromstring(source_xml_content)
        except Exception as e:
            print(f"Error parsing source document XML: {str(e)}")
            return False

        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        
        # ===== SOURCE CELL VALIDATION =====
        source_tables = source_tree.xpath('//w:tbl', namespaces=ns)
        if source_table_index >= len(source_tables) or source_table_index < 0:
            print(f"Error: Source table index {source_table_index} out of range (0-{len(source_tables)-1})")
            return False

        source_rows = source_tables[source_table_index].xpath('.//w:tr', namespaces=ns)
        if source_row_index >= len(source_rows) or source_row_index < 0:
            print(f"Error: Source row index {source_row_index} out of range (0-{len(source_rows)-1})")
            return False

        source_cells = source_rows[source_row_index].xpath('.//w:tc', namespaces=ns)
        if source_col_index >= len(source_cells) or source_col_index < 0:
            print(f"Error: Source column index {source_col_index} out of range (0-{len(source_cells)-1})")
            return False

        source_cell = source_cells[source_col_index]
        source_paragraphs = source_cell.xpath('.//w:p', namespaces=ns)
        
        # ===== CONTENT COPYING =====
        target_cell.text = ""  # Clear target cell


        first_paragraph = True
        n = 0
        pre_left = ''
        pre_hanging = ''
        for i, paragraph in enumerate(source_paragraphs):
            
            # Create or get paragraph
            if first_paragraph:
                new_paragraph = target_cell.paragraphs[0]
                first_paragraph = False
            else:
                new_paragraph = target_cell.add_paragraph()
                    
            # Copy paragraph formatting
            ppr = paragraph.find('.//w:pPr', namespaces=ns)
            if ppr is not None:
                new_ppr = new_paragraph._element.get_or_add_pPr()
                
                numpr = ppr.find('.//w:numPr', namespaces=ns)
                if numpr is not None:
                    try:
                        ilvl = numpr.find('.//w:ilvl', namespaces=ns)
                        numId = numpr.find('.//w:numId', namespaces=ns)
                        if ilvl is not None and numId is not None:
                            # new_paragraph.paragraph_format.left_indent = Pt(18 * int(ilvl.get(qn('w:val'))) + 35)
                            new_numpr = etree.SubElement(new_ppr, qn('w:numPr'))
                            etree.SubElement(new_numpr, qn('w:ilvl'), {qn('w:val'): ilvl.get(qn('w:val'))})
                            etree.SubElement(new_numpr, qn('w:numId'), {qn('w:val'): '1'})
                    except Exception as e:
                        print(f"Warning: Could not copy numbering - {str(e)}")
                
                rpr = ppr.find('.//w:rPr', namespaces=ns)
                if rpr is not None:
                    color = rpr.find('.//w:color', namespaces=ns)
                    if color is not None:
                        if not new_paragraph.runs:
                            new_paragraph.add_run('')
                        
                        for run in new_paragraph.runs:
                            rPr = run._r.get_or_add_rPr()
                            etree.SubElement(rPr, qn('w:color'), {qn('w:val'): color.get(qn('w:val'))})
                
                # Copy the pStyle if it exists
                pstyle = ppr.find('.//w:pStyle', namespaces=ns)
                if pstyle is not None:
                    
                    new_style = etree.SubElement(new_ppr, qn('w:pStyle'))
                    style_name = pstyle.get(qn('w:val'))
                    new_style.set(qn('w:val'), style_name)
                    
                # indent handling
                ind = ppr.find('.//w:ind', namespaces=ns)
                if ind is not None and numpr is None:
                    new_ind = etree.SubElement(new_ppr, qn('w:ind'))
                    
                    left = ind.get(qn('w:left'))
                    # print(type(left))
                    if left is not None:
                        new_ind.set(qn('w:left'), left)
                        pre_left = left
                        n = i
                    hanging = ind.get(qn('w:hanging'))
                    if hanging is not None:
                        new_ind.set(qn('w:hanging'), hanging)
                        pre_hanging = hanging
                    
                
                if i == n + 1 and numpr is None:
                    # Create new indentation element
                    new_ind = etree.SubElement(new_ppr, qn('w:ind'))
                    if pre_left is not None:
                        new_ind.set(qn('w:left'), pre_left)
                    if pre_hanging is not None:
                        new_ind.set(qn('w:hanging'), pre_hanging)
                
                # Spacing handling
                spacing = ppr.find('.//w:spacing', namespaces=ns)
                is_last_paragraph = (i == len(source_paragraphs) - 1)
                is_empty = not paragraph.xpath('.//w:r/w:t', namespaces=ns)

                if spacing is not None:
                    # Space before
                    before = spacing.get(qn('w:before'))
                    if before is not None:
                        new_paragraph.paragraph_format.space_before = Pt(min(int(before)/20, 6))  # Max 6pt
                    
                    # Space after (none for last paragraph)
                    after = spacing.get(qn('w:after'))
                    if after is not None and not is_last_paragraph:
                        new_paragraph.paragraph_format.space_after = Pt(min(int(after)/20, 6))  # Max 6pt
                    else:
                        new_paragraph.paragraph_format.space_after = Pt(0)

                    # Line spacing
                    line = spacing.get(qn('w:line'))
                    if line is not None:
                        line_rule = spacing.get(qn('w:lineRule'), 'auto')
                        line_val = int(line)
                        if line_rule == 'exact':
                            new_paragraph.paragraph_format.line_spacing = Pt(line_val/20)
                        elif line_rule == 'atLeast':
                            new_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                            new_paragraph.paragraph_format.line_spacing = Pt(line_val/20)
                        else:  # auto or multiple
                            new_paragraph.paragraph_format.line_spacing = line_val/240  # 240 = 12pt*20
                else:
                    new_paragraph.paragraph_format.space_after = Pt(0)

            # Hyperlink detection and processing
            is_hyperlink = False
            hyperlink_url = None
            hyperlink_text_runs = []
            hyperlink_style = None

            runs_in_paragraph = paragraph.xpath('./w:r', namespaces=ns)
            for i, run in enumerate(runs_in_paragraph):
                fld_char_begin = run.find('./w:fldChar[@w:fldCharType="begin"]', namespaces=ns)
                instr_text = run.find('./w:instrText', namespaces=ns)
                fld_char_separate = run.find('./w:fldChar[@w:fldCharType="separate"]', namespaces=ns)
                fld_char_end = run.find('./w:fldChar[@w:fldCharType="end"]', namespaces=ns)
                text_element = run.find('./w:t', namespaces=ns)
                rpr = run.find('./w:rPr', namespaces=ns)
                
                # Check for hyperlink style in runs
                if rpr is not None:
                    rstyle = rpr.find('./w:rStyle', namespaces=ns)
                    if rstyle is not None:
                        hyperlink_style = rstyle.get(qn('w:val'))

                if fld_char_begin is not None:
                    is_hyperlink = True
                    hyperlink_url = None
                    hyperlink_text_runs = []

                elif is_hyperlink and instr_text is not None and 'HYPERLINK' in instr_text.text:
                    # Extract URL from HYPERLINK field instruction
                    match = re.search(r'HYPERLINK\s+"([^"]*)"', instr_text.text)
                    if match:
                        hyperlink_url = match.group(1).strip()

                elif is_hyperlink and fld_char_separate is not None:
                    # The actual hyperlink text will come after the separate tag
                    continue

                elif is_hyperlink and text_element is not None:
                    hyperlink_text_runs.append(run)

                elif is_hyperlink and fld_char_end is not None:
                # End of hyperlink - process collected runs
                    hyperlink_text = "".join([t.text for r in hyperlink_text_runs 
                                            if (t := r.find('./w:t', namespaces=ns)) is not None])
                    if hyperlink_url and hyperlink_text:
                        # Create begin field char
                        begin_r = OxmlElement('w:r')
                        begin_fldChar = OxmlElement('w:fldChar')
                        begin_fldChar.set(qn('w:fldCharType'), 'begin')
                        begin_r.append(begin_fldChar)
                        new_paragraph._p.append(begin_r)
                        
                        # Create instrText run
                        instr_r = OxmlElement('w:r')
                        instr_text = OxmlElement('w:instrText')
                        instr_text.set(qn('xml:space'), 'preserve')
                        instr_text.text = f' HYPERLINK "{hyperlink_url}" '
                        instr_r.append(instr_text)
                        new_paragraph._p.append(instr_r)
                        
                        # Create separate field char
                        separate_r = OxmlElement('w:r')
                        separate_fldChar = OxmlElement('w:fldChar')
                        separate_fldChar.set(qn('w:fldCharType'), 'separate')
                        separate_r.append(separate_fldChar)
                        new_paragraph._p.append(separate_r)
                        
                        # Create hyperlink text run with style
                        text_r = OxmlElement('w:r')
                        rPr = OxmlElement('w:rPr')
                        
                        # Add hyperlink style if exists
                        if hyperlink_style:
                            rStyle = OxmlElement('w:rStyle')
                            rStyle.set(qn('w:val'), hyperlink_style)
                            rPr.append(rStyle)
                        
                        # Add formatting elements
                        font_size = OxmlElement('w:sz')
                        font_size.set(qn('w:val'), '22')  # 11pt = 22 half-points
                        rPr.append(font_size)
                        
                        color = OxmlElement('w:color')
                        color.set(qn('w:val'), '0000FF')  # Blue
                        rPr.append(color)
                        
                        underline = OxmlElement('w:u')
                        underline.set(qn('w:val'), 'single')
                        rPr.append(underline)
                        
                        text_r.append(rPr)
                        t = OxmlElement('w:t')
                        t.text = hyperlink_text
                        text_r.append(t)
                        new_paragraph._p.append(text_r)
                        
                        # Create end field char with same formatting
                        end_r = OxmlElement('w:r')
                        end_rPr = OxmlElement('w:rPr')
                        
                        if hyperlink_style:
                            end_rStyle = OxmlElement('w:rStyle')
                            end_rStyle.set(qn('w:val'), hyperlink_style)
                            end_rPr.append(end_rStyle)
                        
                        # Recreate formatting elements for end run
                        end_font_size = OxmlElement('w:sz')
                        end_font_size.set(qn('w:val'), '22')
                        end_rPr.append(end_font_size)
                        
                        end_color = OxmlElement('w:color')
                        end_color.set(qn('w:val'), '0000FF')
                        end_rPr.append(end_color)
                        
                        end_underline = OxmlElement('w:u')
                        end_underline.set(qn('w:val'), 'single')
                        end_rPr.append(end_underline)
                        
                        end_r.append(end_rPr)
                        end_fldChar = OxmlElement('w:fldChar')
                        end_fldChar.set(qn('w:fldCharType'), 'end')
                        end_r.append(end_fldChar)
                        new_paragraph._p.append(end_r)

                elif not is_hyperlink and text_element is not None:
                    text = text_element.text

                    new_run = new_paragraph.add_run(text)
                    
                    # Font name and size
                    new_run.font.name = "Times New Roman"
                    new_run.font.size = Pt(11)
                    # Copy run formatting
                    rpr = run.find('.//w:rPr', namespaces=ns)
                    if rpr is not None:
                        
                        # Color
                        color_element = rpr.find('.//w:color', namespaces=ns)
                        if color_element is not None and color_element.get(qn('w:val')):
                            color_hex = color_element.get(qn('w:val'))
                            try:
                                new_run.font.color.rgb = RGBColor.from_string(color_hex)
                            except ValueError:
                                pass

                        # Bold and italic
                        if rpr.find('.//w:b', namespaces=ns) is not None:
                            new_run.bold = True
                        if rpr.find('.//w:i', namespaces=ns) is not None:
                            new_run.italic = True
                            
                        underline_element = rpr.find('.//w:u', namespaces=ns)
                        if underline_element is not None:
                            underline_val = underline_element.get(qn('w:val'))
                            # Map Word's underline values to WD_UNDERLINE
                            underline_mapping = {
                                'single': WD_UNDERLINE.SINGLE,
                                'double': WD_UNDERLINE.DOUBLE,
                                'thick': WD_UNDERLINE.THICK,
                                'dotted': WD_UNDERLINE.DOTTED,
                                'dash': WD_UNDERLINE.DASH,
                                'dotDash': WD_UNDERLINE.DOT_DASH,
                                'dotDotDash': WD_UNDERLINE.DOT_DOT_DASH,
                                'wave': WD_UNDERLINE.WAVY,
                                'none': WD_UNDERLINE.NONE,
                            }
                            # Default to SINGLE if not specified or invalid
                            underline_style = underline_mapping.get(underline_val, WD_UNDERLINE.SINGLE)
                            new_run.font.underline = underline_style

                            # Underline color (if specified)
                            underline_color = underline_element.get(qn('w:color'))
                            if underline_color:
                                try:
                                    new_run.font.underline_color.rgb = RGBColor.from_string(underline_color)
                                except ValueError:
                                    pass
        return True

    except Exception as e:
        print(f"Error: {e}")
        return False

def copy_row_formatting(source_row, target_row):
    namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    
    for source_cell, target_cell in zip(source_row.cells, target_row.cells):
        source_shading = source_cell._element.find('.//w:shd', namespaces)
        
        if source_shading is not None:
            target_cell_properties = target_cell._element.get_or_add_tcPr()
            target_shading = target_cell_properties.find('.//w:shd', namespaces)
            if target_shading is None:
                target_shading = etree.SubElement(target_cell_properties, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
            
            target_shading.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', 
                               source_shading.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill'))

        source_borders = source_cell._element.find('.//w:tcBorders', namespaces)
        if source_borders is not None:
            target_cell_properties = target_cell._element.get_or_add_tcPr()
            target_borders = target_cell_properties.find('.//w:tcBorders', namespaces)
            if target_borders is None:
                target_borders = etree.SubElement(target_cell_properties, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders')

            for border_element in source_borders.getchildren():
                tag = border_element.tag
                existing_target_border = target_borders.find(tag)
                if existing_target_border is not None:
                    target_borders.remove(existing_target_border)
                target_borders.append(etree.fromstring(etree.tostring(border_element)))

def get_cell_background_color(cell):
    cell_xml = cell._element
    shading = cell_xml.find('.//w:shd', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
    
    if shading is not None:
        fill_color = shading.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
        return fill_color
    return None

def extract_specific_dropdown_pre_display_text(docx_file, dropdown_index):
    try:
        with zipfile.ZipFile(docx_file) as z:
            xml_content = z.read('word/document.xml')
        tree = etree.fromstring(xml_content)
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

        dropdowns = tree.xpath('//w:sdt[.//w:dropDownList]', namespaces=ns)

        if dropdown_index < 0 or dropdown_index >= len(dropdowns):
            return None

        target_dropdown = dropdowns[dropdown_index]
        text_elements = target_dropdown.xpath('.//w:sdtContent//w:t', namespaces=ns)
        pre_display_text = "".join([t.text for t in text_elements]) if text_elements else None

        return pre_display_text

    except Exception as e:
        print(f"Error: {e}")
        return None
    
def find_tables_with_specific_string(doc, search_string):
    matching_table_indexes = []

    for i, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                if search_string in cell.text.strip():
                    matching_table_indexes.append(i)
                    break
            else:
                continue
            break

    return matching_table_indexes

def delete_table_with_paragraphs(doc, table_index, num_paragraphs_above=0, num_paragraphs_below=0):
    try:
        table = doc.tables[table_index]
        table_xml = table._element
        body = doc.element.body

        paragraphs_above_xml = []
        element = table_xml.getprevious()
        count = 0
        while element is not None and count < num_paragraphs_above:
            if element.tag.endswith('p'):
                paragraphs_above_xml.insert(0, element)
                count += 1
            element = element.getprevious()

        paragraphs_below_xml = []
        element = table_xml.getnext()
        count = 0
        while element is not None and count < num_paragraphs_below:
            if element.tag.endswith('p'):
                paragraphs_below_xml.append(element)
                count += 1
            element = element.getnext()

        elements_to_remove = paragraphs_above_xml + [table_xml] + paragraphs_below_xml
        for element in elements_to_remove:
            if element is not None:
                body.remove(element)

        doc.part._element = doc.element

    except IndexError:
        print("Error: Table index out of range.")
    except Exception as e:
        print(f"Error: {e}")
        
def copy_table_with_paragraphs(source_doc, source_table_index, target_doc, target_index, num_paragraphs_above=0, num_paragraphs_below=0):
    try:
        source_table = source_doc.tables[source_table_index]
        table_xml = source_table._element

        paragraphs_above_xml = []
        element = table_xml.getprevious()
        count = 0
        while element is not None and count < num_paragraphs_above:
            if element.tag.endswith('p'):
                paragraphs_above_xml.insert(0, element)
                count += 1
            element = element.getprevious()

        paragraphs_below_xml = []
        element = table_xml.getnext()
        count = 0
        while element is not None and count < num_paragraphs_below:
            if element.tag.endswith('p'):
                paragraphs_below_xml.append(element)
                count += 1
            element = element.getnext()

        new_table_xml = etree.fromstring(etree.tostring(table_xml))

        target_body = target_doc.element.body

        if target_index < len(target_doc.tables):
            target_table_xml = target_doc.tables[target_index]._element
            insert_before = target_table_xml
        else:
            insert_before = None

        if insert_before is not None:
            for p_xml in paragraphs_above_xml:
                target_body.insert(target_body.index(insert_before), p_xml)
            target_body.insert(target_body.index(insert_before), new_table_xml)
            for p_xml in paragraphs_below_xml:
                target_body.insert(target_body.index(insert_before), p_xml)
        else:
            for p_xml in paragraphs_above_xml:
                target_body.append(p_xml)
            target_body.append(new_table_xml)
            for p_xml in paragraphs_below_xml:
                target_body.append(p_xml)

        target_doc.part._element = target_doc.element

    except IndexError:
        print("Error: Table index out of range.")
    except Exception as e:
        print(f"Error: {e}")
        
def add_page_break_before_table(doc, table_index):
    try:
        table = doc.tables[table_index]
        table_element = table._element
        body = doc.element.body

        new_paragraph = etree.Element(qn('w:p'))
        new_run = etree.SubElement(new_paragraph, qn('w:r'))
        new_br = etree.SubElement(new_run, qn('w:br'), {qn('w:type'): 'page'})

        body.insert(body.index(table_element), new_paragraph)

        doc.part._element = doc.element

    except IndexError:
        print("Error: Table index out of range.")
    except Exception as e:
        print(f"Error: {e}")

def write_text_to_cell(doc, table_index, row_index, cell_index, text, font_size, bold=True, alignment="left"):
    try:
        table = doc.tables[table_index]
        cell = table.rows[row_index].cells[cell_index]
        paragraph = cell.paragraphs[0]
        run = paragraph.clear().add_run(text)

        font = run.font
        font.name = "Times New Roman"
        font.size = Pt(font_size)
        font.bold = bold

        if alignment == "left":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif alignment == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == "right":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            print("Warning: Invalid alignment specified. Defaulting to left.")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    except IndexError:
        print("Error: Table, row, or cell index out of range.")
    except Exception as e:
        print(f"Error: {e}")
        
def copy_publishable_matter_to_target(source_doc, target_doc, source_table_index, target_table_index):
    source_publishableMatter_table = source_doc.tables[source_table_index]
    target_publishableMatter_table = target_doc.tables[target_table_index]
    
    copy_cell_content_to_target_cell(source_doc, source_table_index, 2, 0, target_doc, target_table_index, 3, 0)
    copy_cell_content_to_target_cell(source_doc, source_table_index, 4, 0, target_doc, target_table_index, 5, 0)
    copy_cell_content_to_target_cell(source_doc, source_table_index, 5, 1, target_doc, target_table_index, 7, 0)
    
    for i in range(len(source_publishableMatter_table.rows)):
        cell_text = source_publishableMatter_table.cell(i, 0).text.strip()
        if "Was it cross-border?" in cell_text:
            crossBorder_text = ""
            temp_color = get_cell_background_color(source_publishableMatter_table.cell(i + 2, 0))
            j = 3
            if temp_color != "D0CECE":
                crossBorder_text = source_publishableMatter_table.cell(i + 2, 0).text.strip()

            while True:
                try:
                    current_color = get_cell_background_color(source_publishableMatter_table.cell(i + j, 0))
                    if current_color == temp_color and current_color != "D0CECE" and source_publishableMatter_table.cell(i + j, 0).text.strip() != "":
                        crossBorder_text += ", " + source_publishableMatter_table.cell(i + j, 0).text.strip()
                        j += 1
                    else:
                        break
                except IndexError:
                    break

            write_text_to_cell(target_doc, target_table_index, 9, 0, crossBorder_text, 11, bold=False, alignment="left")
            
        keywords = ["Lead partner", "Lead Consultant solicitor", "Lead counsel", "Lead director", "Lead associate"]
        if any(keyword in cell_text for keyword in keywords):
            j = 3
            leadPartner_text = ""
            temp_color = get_cell_background_color(source_publishableMatter_table.cell(i + 2, 0))
            if temp_color != "D0CECE":
                leadPartner_text = source_publishableMatter_table.cell(i + 2, 0).text.strip()
                if "Practice area" in source_publishableMatter_table.cell(i + 1, 4).text.strip():
                    if source_publishableMatter_table.cell(i + 2, 4).text.strip() != "":
                        leadPartner_text = leadPartner_text + " - " + source_publishableMatter_table.cell(i + 2, 4).text.strip()
                else:
                    if source_publishableMatter_table.cell(i + 2, 5).text.strip() != "":
                        leadPartner_text = leadPartner_text + " - " + source_publishableMatter_table.cell(i + 2, 5).text.strip()
            while True:
                try:
                    current_color = get_cell_background_color(source_publishableMatter_table.cell(i + j, 0))
                    if current_color != "D0CECE" and source_publishableMatter_table.cell(i + j, 0).text.strip() != "":
                        if "Practice area" in source_publishableMatter_table.cell(i + 1, 4).text.strip():
                            if source_publishableMatter_table.cell(i + 2, 4).text.strip() != "":
                                leadPartner_text += ", " + source_publishableMatter_table.cell(i + j, 0).text.strip() + " - " + source_publishableMatter_table.cell(i + j, 4).text.strip()
                        else:
                            if source_publishableMatter_table.cell(i + 2, 5).text.strip() != "":
                                leadPartner_text += ", " + source_publishableMatter_table.cell(i + j, 0).text.strip() + " - " + source_publishableMatter_table.cell(i + j, 5).text.strip()
                        j += 1
                    else:
                        break
                except IndexError:
                    break

            write_text_to_cell(target_doc, target_table_index, 11, 0, leadPartner_text, 11, bold=False, alignment="left")
        
        if "Other key team members" in cell_text:
            j = 3
            otherKeymembers_text = ""
            temp_color = get_cell_background_color(source_publishableMatter_table.cell(i + 2, 0))
            if temp_color != "D0CECE":
                otherKeymembers_text = source_publishableMatter_table.cell(i + 2, 0).text.strip()
                if "Practice area" in source_publishableMatter_table.cell(i + 1, 4).text.strip():
                    if source_publishableMatter_table.cell(i + 2, 4).text.strip() != "":
                        otherKeymembers_text = otherKeymembers_text + " - " + source_publishableMatter_table.cell(i + 2, 4).text.strip()
                else:
                    if source_publishableMatter_table.cell(i + 2, 5).text.strip() != "":
                        otherKeymembers_text = otherKeymembers_text + " - " + source_publishableMatter_table.cell(i + 2, 5).text.strip()

            while True:
                try:
                    current_color = get_cell_background_color(source_publishableMatter_table.cell(i + j, 0))
                    if current_color != "D0CECE" and source_publishableMatter_table.cell(i + j, 0).text.strip() != "":
                        if "Practice area" in source_publishableMatter_table.cell(i + 1, 4).text.strip():
                            if source_publishableMatter_table.cell(i + 2, 4).text.strip() != "":
                                otherKeymembers_text += ", " + source_publishableMatter_table.cell(i + j, 0).text.strip() + " - " + source_publishableMatter_table.cell(i + j, 4).text.strip()
                        else:
                            if source_publishableMatter_table.cell(i + 2, 5).text.strip() != "":
                                otherKeymembers_text += ", " + source_publishableMatter_table.cell(i + j, 0).text.strip() + " - " + source_publishableMatter_table.cell(i + j, 5).text.strip()
                        j += 1
                    else:
                        break
                except IndexError:
                    break

            write_text_to_cell(target_doc, target_table_index, 13, 0, otherKeymembers_text, 11, bold=False, alignment="left")
            
        if "Other firms advising on the matter and their role(s)" in cell_text:
            j = 3
            otherFirmsadvising_text = ""
            temp_color = get_cell_background_color(source_publishableMatter_table.cell(i + 2, 0))
            if temp_color != "D0CECE":
                otherFirmsadvising_text = source_publishableMatter_table.cell(i + 2, 0).text.strip()
                
                if "Advising" in source_publishableMatter_table.cell(i + 1, 4).text.strip():
                    if source_publishableMatter_table.cell(i + 2, 2).text.strip() != "":
                        otherFirmsadvising_text = otherFirmsadvising_text + " - " + source_publishableMatter_table.cell(i + 2, 2).text.strip()
                    if source_publishableMatter_table.cell(i + 2, 4).text.strip() != "":
                        otherFirmsadvising_text = otherFirmsadvising_text + " - " + source_publishableMatter_table.cell(i + 2, 4).text.strip()
                else:
                    if source_publishableMatter_table.cell(i + 2, 3).text.strip() != "":
                        otherFirmsadvising_text = otherFirmsadvising_text + " - " + source_publishableMatter_table.cell(i + 2, 3).text.strip()
                    if source_publishableMatter_table.cell(i + 2, 5).text.strip() != "":
                        otherFirmsadvising_text = otherFirmsadvising_text + " - " + source_publishableMatter_table.cell(i + 2, 5).text.strip()
            while True:
                try:
                    current_color = get_cell_background_color(source_publishableMatter_table.cell(i + j, 0))
                    if current_color != "D0CECE" and source_publishableMatter_table.cell(i + j, 0).text.strip() != "":
                        otherFirmsadvising_text = otherFirmsadvising_text + "; " + source_publishableMatter_table.cell(i + j, 0).text.strip()
                        
                        if "Advising" in source_publishableMatter_table.cell(i + 1, 4).text.strip():
                            if source_publishableMatter_table.cell(i + j, 2).text.strip() != "":
                                otherFirmsadvising_text = otherFirmsadvising_text + " - " + source_publishableMatter_table.cell(i + j, 2).text.strip()
                            if source_publishableMatter_table.cell(i + j, 4).text.strip() != "":
                                otherFirmsadvising_text = otherFirmsadvising_text + " - " + source_publishableMatter_table.cell(i + j, 4).text.strip()
                        else:
                            if source_publishableMatter_table.cell(i + j, 3).text.strip() != "":
                                otherFirmsadvising_text = otherFirmsadvising_text + " - " + source_publishableMatter_table.cell(i + j, 3).text.strip()
                            if source_publishableMatter_table.cell(i + j, 5).text.strip() != "":
                                otherFirmsadvising_text = otherFirmsadvising_text + " - " + source_publishableMatter_table.cell(i + j, 5).text.strip()
                        j += 1
                    else:
                        break
                except IndexError:
                    break

            write_text_to_cell(target_doc, target_table_index, 15, 0, otherFirmsadvising_text, 11, bold=False, alignment="left")
            
        if "Start date" in cell_text:
            write_text_to_cell(target_doc, target_table_index, 17, 0, source_publishableMatter_table.cell(i + 1, 4).text.strip(), 11, bold=False, alignment="left")
        
async def l500_chamber_convert(source_docx_path, target_docx_path):
    # Load the source and target documents
    source_doc = Document(source_docx_path)
    target_doc = Document(target_docx_path)

    # Extract the Firm Name
    copy_cell_content_to_target_cell(source_doc, 0, 0, 0,
                                      target_doc, 0, 1, 0)
    
    # Extract the Practice Area
    practiceArea_text = extract_specific_dropdown_pre_display_text(source_docx_path, 0)
    write_text_to_cell(target_doc, 1, 1, 0, practiceArea_text, 11, bold=False, alignment="left")
    
    # Extract the Location
    location_text = extract_specific_textbox_text(source_docx_path, 0)
    write_text_to_cell(target_doc, 2, 1, 0, location_text, 11, bold=False, alignment="left")
    
    # Extract the Contact Details
    source_contactDetail_table = source_doc.tables[1]
    target_contactDetail_table = target_doc.tables[3]
    for i in range(len(target_contactDetail_table.rows) - 1, 2, -1):
        temp_row = target_contactDetail_table.rows[i]
        target_contactDetail_table._element.remove(temp_row._element)
    for i in range(1, len(source_contactDetail_table.rows)):
        if source_contactDetail_table.cell(i, 0).text.strip() != "":  
            if i < len(source_contactDetail_table.rows) - 1:
                source_row = target_contactDetail_table.rows[2]
                new_row = target_contactDetail_table.add_row()
                copy_row_formatting(source_row, new_row)
            copy_cell_content_to_target_cell(source_doc, 1, i, 0, target_doc, 3, i + 1, 0)
            copy_cell_content_to_target_cell(source_doc, 1, i, 2, target_doc, 3, i + 1, 1)
            copy_cell_content_to_target_cell(source_doc, 1, i, 3, target_doc, 3, i + 1, 2)
            
    
    # Extract the Department Name    
    copy_cell_content_to_target_cell(source_doc, 2, 0, 0, target_doc, 4, 1, 0)
    
    # Extract the Department Information
    copy_cell_content_to_target_cell(source_doc, 4, 0, 1, target_doc, 5, 1, 0)
    copy_cell_content_to_target_cell(source_doc, 4, 0, 4, target_doc, 5, 3, 0)
    
    # Extract the Heads of Team(Department)
    source_headsOfteam_table = source_doc.tables[3]
    target_headsOfteam_table = target_doc.tables[6]
    for i in range(len(target_headsOfteam_table.rows) - 1, 2, -1):
        temp_row = target_headsOfteam_table.rows[i]
        target_headsOfteam_table._element.remove(temp_row._element)
    for i in range(1, len(source_headsOfteam_table.rows)):
        if source_headsOfteam_table.cell(i, 0).text.strip() != "":   
            if i < len(source_headsOfteam_table.rows) - 1:
                source_row = target_headsOfteam_table.rows[2]
                new_row = target_headsOfteam_table.add_row()
                copy_row_formatting(source_row, new_row)
            copy_cell_content_to_target_cell(source_doc, 3, i, 0, target_doc, 6, i + 1, 0)
        
    # Extract the factor of best department
    copy_cell_content_to_target_cell(source_doc, 5, 0, 0, target_doc, 9, 1, 0)
    
    # Extract the client's feedback
    feedback_indices = find_tables_with_specific_string(source_doc, search_string="Comments")
    if len(feedback_indices) > 0:
        source_feedback_table_index = feedback_indices[0]
        # print(source_feedback_table_index)
        source_feedback_table = source_doc.tables[source_feedback_table_index]
        target_feedback_table = target_doc.tables[10]
        for i in range(len(target_feedback_table.rows) - 1, 2, -1):
            temp_row = target_feedback_table.rows[i]
            target_feedback_table._element.remove(temp_row._element)
        for i in range(1, len(source_feedback_table.rows)):
            # print(source_feedback_table.cell(i, 0).text.strip())
            if source_feedback_table.cell(i, 0).text.strip() != "":
                source_row = target_feedback_table.rows[2]
                new_row = target_feedback_table.add_row()
                copy_row_formatting(source_row, new_row)
                copy_cell_content_to_target_cell(source_doc, source_feedback_table_index, i, 0, target_doc, 10, i + 1, 0)
                copy_cell_content_to_target_cell(source_doc, source_feedback_table_index, i, 1, target_doc, 10, i + 1, 1)
                copy_cell_content_to_target_cell(source_doc, source_feedback_table_index, i, 3, target_doc, 10, i + 1, 2)
    
    # Extract feedback in this practice area
    copy_cell_content_to_target_cell(source_doc, 7, 0, 0, target_doc, 11, 1, 0)
    
    # Extract the publishable clients
    source_publishableClients_table = source_doc.tables[8]
    target_publishableClients_table = target_doc.tables[12]
    for i in range(len(target_publishableClients_table.rows) - 1, 1, -1):
        temp_row = target_publishableClients_table.rows[i]
        target_publishableClients_table._element.remove(temp_row._element)
    for i in range(1, len(source_publishableClients_table.rows) - 1):
        if source_publishableClients_table.cell(i, 0).text.strip() != "":
            source_row = target_publishableClients_table.rows[1]
            new_row = target_publishableClients_table.add_row()
            copy_row_formatting(source_row, new_row)
            target_publishableClients_table.cell(i + 1, 0).text = str(i)
            copy_cell_content_to_target_cell(source_doc, 8, i, 0, target_doc, 12, i + 1, 1)
            copy_cell_content_to_target_cell(source_doc, 8, i, 1, target_doc, 12, i + 1, 2)
    last_publish_index = len(source_publishableClients_table.rows) - 1
    last__publish_text = source_publishableClients_table.cell(last_publish_index, 0).text.strip()
    # print(last__publish_text)
    if "To add more clients, right-click in any field and select" not in last__publish_text and last__publish_text != "":
        source_row = target_publishableClients_table.rows[1]
        new_row = target_publishableClients_table.add_row()
        copy_row_formatting(source_row, new_row)
        target_publishableClients_table.cell(last_publish_index + 1, 0).text = str(last_publish_index)
        copy_cell_content_to_target_cell(source_doc, 8, last_publish_index, 0, target_doc, 12, last_publish_index + 1, 1)
        copy_cell_content_to_target_cell(source_doc, 8, last_publish_index, 1, target_doc, 12, last_publish_index + 1, 2)
    
    # Extract Confidential clients
    source_confidentialClients_table = source_doc.tables[9]
    confidentialClients_table_indexlist = find_tables_with_specific_string(target_doc, search_string="E0 – CONFIDENTIAL CLIENTS")
    confidentialClients_table_index = confidentialClients_table_indexlist[0]
    target_confidentialClients_table = target_doc.tables[confidentialClients_table_index]
    for i in range(len(target_confidentialClients_table.rows) - 1, 1, -1):
        temp_row = target_confidentialClients_table.rows[i]
        target_confidentialClients_table._element.remove(temp_row._element)
    for i in range(1, len(source_confidentialClients_table.rows) - 1):
        if source_confidentialClients_table.cell(i, 0).text.strip() != "":
            source_row = target_confidentialClients_table.rows[1]
            new_row = target_confidentialClients_table.add_row()
            copy_row_formatting(source_row, new_row)
            target_confidentialClients_table.cell(i + 1, 0).text = str(i)
            copy_cell_content_to_target_cell(source_doc, 9, i, 0, target_doc, confidentialClients_table_index, i + 1, 1)
            copy_cell_content_to_target_cell(source_doc, 9, i, 1, target_doc, confidentialClients_table_index, i + 1, 2)
    last_confidential_index = len(source_confidentialClients_table.rows) - 1
    last__confidential_text = source_confidentialClients_table.cell(last_confidential_index, 0).text.strip()
    # print(last__confidential_text)
    if "To add more clients, right-click in any field and select" not in last__confidential_text and last__confidential_text != "":
        source_row = target_confidentialClients_table.rows[1]
        new_row = target_confidentialClients_table.add_row()
        copy_row_formatting(source_row, new_row)
        target_confidentialClients_table.cell(i + 1, 0).text = str(last_confidential_index)
        copy_cell_content_to_target_cell(source_doc, 8, last_confidential_index, 0, target_doc, 12, last_confidential_index + 1, 1)
        copy_cell_content_to_target_cell(source_doc, 8, last_confidential_index, 1, target_doc, 12, last_confidential_index + 1, 2)
        
    # Extract the information of Ranked and Unranked lawyers
    leadingPartner_indices = []
    leadingAssociate_indices = []
    nextGenerationPartner_indices = []
    leadingPartner_indices = find_tables_with_specific_string(source_doc, search_string="Partner: leading partner")
    leadingPartner_indices += find_tables_with_specific_string(source_doc, search_string="Partner: leading individual")
    nextGenerationPartner_indices = find_tables_with_specific_string(source_doc, search_string="Partner: next generation")
    leadingAssociate_indices = find_tables_with_specific_string(source_doc, search_string="Associate: leading associate")
    leadingAssociate_indices += find_tables_with_specific_string(source_doc, search_string="Associate: rising star")
    # print(leadingAssociate_indices)
    ranked_unrankedLawyers_indices = leadingPartner_indices + nextGenerationPartner_indices
    # print(ranked_unrankedLawyers_indices)
    target_ranked_unrankedLawyers_table = target_doc.tables[8]
    for i in range(len(target_ranked_unrankedLawyers_table.rows) - 1, 2, -1):
        temp_row = target_ranked_unrankedLawyers_table.rows[i]
        target_ranked_unrankedLawyers_table._element.remove(temp_row._element)
    ranked_unrankedLawyers_indices_length = len(ranked_unrankedLawyers_indices)
    leadingAssociate_indices_length = len(leadingAssociate_indices)
    
    if ranked_unrankedLawyers_indices_length > 0:
        for i in range(1, ranked_unrankedLawyers_indices_length):
            temp_index = ranked_unrankedLawyers_indices[i]
            source_ranked_unrankedLawyers_table = source_doc.tables[temp_index]
            if source_ranked_unrankedLawyers_table.cell(2, 0).text.strip != "":
                source_row = target_ranked_unrankedLawyers_table.rows[2]
                new_row = target_ranked_unrankedLawyers_table.add_row()
                copy_row_formatting(source_row, new_row)
                
                copy_cell_content_to_target_cell(source_doc, temp_index, 2, 0, target_doc, 8, i + 2, 0)
                copy_cell_content_to_target_cell(source_doc, temp_index, 4, 0, target_doc, 8, i + 2, 1)
                write_text_to_cell(target_doc, 8, i + 2, 2, "Y", 11, bold=False, alignment="left")
    if leadingAssociate_indices_length > 0:        
        for i in range(leadingAssociate_indices_length):
            temp_index = leadingAssociate_indices[i]
            source_ranked_unrankedLawyers_table = source_doc.tables[temp_index]
            if source_ranked_unrankedLawyers_table.cell(2, 0).text.strip != "":
                source_row = target_ranked_unrankedLawyers_table.rows[2]
                new_row = target_ranked_unrankedLawyers_table.add_row()
                copy_row_formatting(source_row, new_row)
                
                copy_cell_content_to_target_cell(source_doc, temp_index, 2, 0, target_doc, 8, i + ranked_unrankedLawyers_indices_length + 2, 0)
                copy_cell_content_to_target_cell(source_doc, temp_index, 4, 0, target_doc, 8, i + ranked_unrankedLawyers_indices_length + 2, 1)
                write_text_to_cell(target_doc, 8, i + ranked_unrankedLawyers_indices_length + 2, 2, "N", 11, bold=False, alignment="left")
    if ranked_unrankedLawyers_indices_length > 0:
        copy_cell_content_to_target_cell(source_doc, ranked_unrankedLawyers_indices[0], 2, 0, target_doc, 8, 2, 0)
        copy_cell_content_to_target_cell(source_doc, ranked_unrankedLawyers_indices[0], 4, 0, target_doc, 8, 2, 1)
        write_text_to_cell(target_doc, 8, 2, 2, "Y", 11, bold=False, alignment="left")
    
    
    # Extract Hires/Departures of partners in last 12 months
    source_hiresDepartures_table_indices = set(find_tables_with_specific_string(source_doc, search_string="Name"))
    temp_indices = set(find_tables_with_specific_string(source_doc, search_string="Position/role"))
    source_hiresDepartures_table_indices = list(source_hiresDepartures_table_indices.intersection(temp_indices))
    source_hiresDepartures_table_indices.sort()
    if len(source_hiresDepartures_table_indices) > 0:
        source_hiresDepartures_table_index = source_hiresDepartures_table_indices[0]
        source_hiresDepartures_table = source_doc.tables[source_hiresDepartures_table_index]
        target_hiresDepartures_table = target_doc.tables[7]
        for i in range(len(target_hiresDepartures_table.rows) - 1, 2, -1):
            row = target_hiresDepartures_table.rows[i]
            target_hiresDepartures_table._element.remove(row._element)
        for i in range(1, len(source_hiresDepartures_table.rows) - 1):
            if source_hiresDepartures_table.cell(i, 0).text.strip() != "":
                source_row = target_hiresDepartures_table.rows[2]
                new_row = target_hiresDepartures_table.add_row()
                copy_row_formatting(source_row, new_row)
                copy_cell_content_to_target_cell(source_doc, source_hiresDepartures_table_index, i, 0, target_doc, 7, i + 1, 0)
                copy_cell_content_to_target_cell(source_doc, source_hiresDepartures_table_index, i, 2, target_doc, 7, i + 1, 1)
                copy_cell_content_to_target_cell(source_doc, source_hiresDepartures_table_index, i, 3, target_doc, 7, i + 1, 2)
            
    # Extract Publishable Matter
    publishable_matter_indices = set(find_tables_with_specific_string(source_doc, search_string="Publishable matter"))
    temp_indices = set(find_tables_with_specific_string(source_doc, search_string="Name of client"))
    publishable_matter_indices = list(publishable_matter_indices.intersection(temp_indices))
    publishable_matter_indices.sort()
    print(publishable_matter_indices)
    
    non_publishable_matter_indices1 = set(find_tables_with_specific_string(source_doc, search_string="Non-publishable matter"))
    non_publishable_matter_indices = list(non_publishable_matter_indices1.intersection(temp_indices))
    non_publishable_matter_indices.sort()
    
    print(non_publishable_matter_indices)
    publishable_matter_list_length = len(publishable_matter_indices)
    non_publishable_matter_list_length = len(non_publishable_matter_indices)
    matter_list_length = publishable_matter_list_length + non_publishable_matter_list_length
    
    confidentialClients_table_indexlist = find_tables_with_specific_string(target_doc, search_string="E0 – CONFIDENTIAL CLIENTS")
    confidentialClients_table_index = confidentialClients_table_indexlist[0]
    
    
    for i in range(9):
        delete_table_with_paragraphs(target_doc, 14, num_paragraphs_above=0, num_paragraphs_below=3)
    for i in range(9):
        delete_table_with_paragraphs(target_doc, 16, num_paragraphs_above=0, num_paragraphs_below=3)
    
    if publishable_matter_list_length == 0 and non_publishable_matter_list_length != 0:
        delete_table_with_paragraphs(target_doc, 13, num_paragraphs_above=0, num_paragraphs_below=3)
        delete_table_with_paragraphs(target_doc, 12, num_paragraphs_above=8, num_paragraphs_below=0)
        for i in range(non_publishable_matter_list_length - 1):
            copy_table_with_paragraphs(target_doc, 13, target_doc, 14, num_paragraphs_above=0, num_paragraphs_below=0)
            
        for i in range(non_publishable_matter_list_length - 2):
            add_page_break_before_table(target_doc, 15 + i)  
            
        for i in range(non_publishable_matter_list_length):
            temp_text1 = ""
            if i == 0:
                temp_text1 = "Confidential Work Highlights in last 12 months"
            temp_text2 = "Confidential Matter " + str(i + 1)
            write_text_to_cell(target_doc, 13 + i, 0, 0, temp_text1, 13, alignment="left")
            write_text_to_cell(target_doc, 13 + i, 1, 0, temp_text2, 14, alignment="center") 
        
        for i in range(non_publishable_matter_list_length):
            copy_publishable_matter_to_target(source_doc, target_doc, non_publishable_matter_indices[i], 13 + i)
    
    elif publishable_matter_list_length != 0 and non_publishable_matter_list_length == 0:
        delete_table_with_paragraphs(target_doc, 15, num_paragraphs_above=0, num_paragraphs_below=3)
        delete_table_with_paragraphs(target_doc, 14, num_paragraphs_above=8, num_paragraphs_below=0)
        for i in range(publishable_matter_list_length - 1):
            copy_table_with_paragraphs(target_doc, 13, target_doc, 13, num_paragraphs_above=0, num_paragraphs_below=0)

        for i in range(publishable_matter_list_length - 1):
            add_page_break_before_table(target_doc, 14 + i)  
        
        for i in range(publishable_matter_list_length):
            temp_text1 = ""
            if i == 0:
                temp_text1 = "Publishable Work Highlights in last 12 months"
            temp_text2 = "Publishable Matter " + str(i + 1)
            write_text_to_cell(target_doc, 13 + i, 0, 0, temp_text1, 13, alignment="left")
            write_text_to_cell(target_doc, 13 + i, 1, 0, temp_text2, 14, alignment="center")
        
        for i in range(publishable_matter_list_length):
            copy_publishable_matter_to_target(source_doc, target_doc, publishable_matter_indices[i], 13 + i)
        
    else:    
        for i in range(publishable_matter_list_length - 1):
            copy_table_with_paragraphs(target_doc, 13, target_doc, 13, num_paragraphs_above=0, num_paragraphs_below=0)

        for i in range(publishable_matter_list_length - 1):
            add_page_break_before_table(target_doc, 14 + i)  
        
        for i in range(publishable_matter_list_length):
            temp_text1 = ""
            if i == 0:
                temp_text1 = "Publishable Work Highlights in last 12 months"
            temp_text2 = "Publishable Matter " + str(i + 1)
            write_text_to_cell(target_doc, 13 + i, 0, 0, temp_text1, 13, alignment="left")
            write_text_to_cell(target_doc, 13 + i, 1, 0, temp_text2, 14, alignment="center")
    
        for i in range(non_publishable_matter_list_length - 1):
            copy_table_with_paragraphs(target_doc, 14 + publishable_matter_list_length, target_doc, 14 + publishable_matter_list_length, num_paragraphs_above=0, num_paragraphs_below=0)
            
        for i in range(non_publishable_matter_list_length - 1):
            add_page_break_before_table(target_doc, 15 + publishable_matter_list_length + i)  
            
        for i in range(non_publishable_matter_list_length):
            temp_text1 = ""
            if i == 0:
                temp_text1 = "Confidential Work Highlights in last 12 months"
            temp_text2 = "Confidential Matter " + str(i + 1)
            write_text_to_cell(target_doc, 14 + publishable_matter_list_length + i, 0, 0, temp_text1, 13, alignment="left")
            write_text_to_cell(target_doc, 14 + publishable_matter_list_length + i, 1, 0, temp_text2, 14, alignment="center")        
     
        for i in range(publishable_matter_list_length):
            copy_publishable_matter_to_target(source_doc, target_doc, publishable_matter_indices[i], 13 + i)
        
        for i in range(non_publishable_matter_list_length):
            copy_publishable_matter_to_target(source_doc, target_doc, non_publishable_matter_indices[i], 14 + publishable_matter_list_length + i)
        
        
    # Save the modified target document
    file_name_without_extension = os.path.splitext(source_docx_path)[0]
    file_name = f"{file_name_without_extension}_result.docx"
    target_doc.save(file_name)
    print("Content copied to the target document successfully.")
    
    return file_name
    


async def validate_document(source_path):
    
    source_doc = Document(source_path)
    publishable_matter_indices = set(find_tables_with_specific_string(source_doc, search_string="Publishable matter"))
    temp_indices = set(find_tables_with_specific_string(source_doc, search_string="Name of client"))
    publishable_matter_indices = list(publishable_matter_indices.intersection(temp_indices))
    publishable_matter_indices.sort()
    publishable_matter_list_length = len(publishable_matter_indices)
    # print(publishable_matter_indices)
    
    non_publishable_matter_indices1 = set(find_tables_with_specific_string(source_doc, search_string="Non-publishable matter"))
    non_publishable_matter_indices = list(non_publishable_matter_indices1.intersection(temp_indices))
    non_publishable_matter_indices.sort()
    non_publishable_matter_list_length = len(non_publishable_matter_indices)

    matter_indices = publishable_matter_indices + non_publishable_matter_indices
    matter_indices.sort()
    
    matter_indices_length = len(matter_indices)
    
    # print(matter_indices)
    for i in range(matter_indices_length):
        matter_table_index = matter_indices[i]
        matter_table = source_doc.tables[matter_table_index]
        matter_table_length = len(matter_table.rows)
        top = 0
        bottom = 0
        for j in range(matter_table_length):
            cell_text = matter_table.cell(j, 0).text.strip()
            if "Publishable matter" in cell_text or "Non-publishable matter" in cell_text:
                top = j
            if "Start date" in cell_text:
                bottom = j
        if bottom + 2 < matter_table_length:
            for k in range(bottom + 2, matter_table_length):
                delete_table_row(source_doc, matter_table_index, bottom + 2)
        if top != 0:
            for k in range(top):
                delete_table_row(source_doc, matter_table_index, 0)
    
    publisahble_matter_table_first_cell_texts = []
    for i in range(publishable_matter_list_length):
        publishable_matter_table_index = matter_indices[i]
        publishable_matter_table = source_doc.tables[publishable_matter_table_index]  
        publisahble_matter_table_first_cell_texts.append(publishable_matter_table.cell(0, 0).text.strip())
        
    publisahble_matter_table_numbers = [int(re.search(r'\d+', s).group()) for s in publisahble_matter_table_first_cell_texts]
    # print(publisahble_matter_table_numbers)
    publisahble_matter_table_missing_numbers = []
    if len(publisahble_matter_table_numbers) > 0:
        publisahble_matter_table_min_num = 1
        publisahble_matter_table_max_num = publisahble_matter_table_numbers[-1]
    
        # Generate the full expected consecutive sequence
        publisahble_matter_table_full_sequence = set(range(publisahble_matter_table_min_num, publisahble_matter_table_max_num + 1))
        publisahble_matter_table_actual_sequence = set(publisahble_matter_table_numbers)
        
        publisahble_matter_table_missing_numbers = sorted(publisahble_matter_table_full_sequence - publisahble_matter_table_actual_sequence)
    
        if publisahble_matter_table_missing_numbers:
            # print(publisahble_matter_table_missing_numbers)
            publishable_str = ', '.join(f"Publishable Matter {num}" for num in publisahble_matter_table_missing_numbers)
            # Construct the final message
            message = f"{publishable_str} was written wrong. Please rewrite it."
            print(message)
            return False
    
    non_publisahble_matter_table_first_cell_texts = []
    for i in range(publishable_matter_list_length, matter_indices_length):
        non_publishable_matter_table_index = matter_indices[i]
        non_publishable_matter_table = source_doc.tables[non_publishable_matter_table_index]  
        non_publisahble_matter_table_first_cell_texts.append(non_publishable_matter_table.cell(0, 0).text.strip())
    # print(non_publisahble_matter_table_first_cell_texts)
    non_publisahble_matter_table_numbers = [int(re.search(r'\d+', s).group()) for s in non_publisahble_matter_table_first_cell_texts]
    # print(non_publisahble_matter_table_numbers)
    non_publisahble_matter_table_missing_numbers = []
    if len(non_publisahble_matter_table_numbers) > 0:
        non_publisahble_matter_table_min_num = 1
        non_publisahble_matter_table_max_num = non_publisahble_matter_table_numbers[-1]
    
        # Generate the full expected consecutive sequence
        non_publisahble_matter_table_full_sequence = set(range(non_publisahble_matter_table_min_num, non_publisahble_matter_table_max_num + 1))
        non_publisahble_matter_table_actual_sequence = set(non_publisahble_matter_table_numbers)
    
        non_publisahble_matter_table_missing_numbers = sorted(non_publisahble_matter_table_full_sequence - non_publisahble_matter_table_actual_sequence)
        # print(non_publisahble_matter_table_missing_numbers)
        if non_publisahble_matter_table_missing_numbers:
            # print(non_publisahble_matter_table_missing_numbers)
            non_publishable_str = ', '.join(f"Non-Publishable Matter {num}" for num in non_publisahble_matter_table_missing_numbers)
            # Construct the final message
            message = f"{non_publishable_str} was written wrong. Please rewrite it."
            print(message)    
            return False
    
    if not publisahble_matter_table_missing_numbers and not non_publisahble_matter_table_missing_numbers:
        leadingPartner_indices = set()

        leadingPartner_indices.update(find_tables_with_specific_string(source_doc, search_string="Partner: leading partner"))
        leadingPartner_indices.update(find_tables_with_specific_string(source_doc, search_string="Partner: leading individual"))
        leadingPartner_indices.update(find_tables_with_specific_string(source_doc, search_string="Partner: next generation partner"))
        leadingPartner_indices.update(find_tables_with_specific_string(source_doc, search_string="Associate: leading associate"))
        leadingPartner_indices.update(find_tables_with_specific_string(source_doc, search_string="Associate: rising star"))

        temp_indices = set(find_tables_with_specific_string(source_doc, search_string="Supporting information"))

        leadingPartner_indices = list(leadingPartner_indices.intersection(temp_indices))
        leadingPartner_indices.sort()
        # print(leadingPartner_indices)
        leadingPartner_indices_length = len(leadingPartner_indices)
        for i in range(leadingPartner_indices_length):
            leadingPartner_table_index = leadingPartner_indices[i]
            leadingPartner_table = source_doc.tables[leadingPartner_table_index]
            leadingPartner_table_length = len(leadingPartner_table.rows)
            top = 0
            for j in range(leadingPartner_table_length):
                cell_text = leadingPartner_table.cell(j, 0).text.strip()
                if any(s in cell_text for s in ["Partner: leading partner", "Partner: leading individual", "Partner: next generation partner", "Associate: leading associate", "Associate: rising star"]):
                    top = j

            if top != 0:
                for k in range(top):
                    delete_table_row(source_doc, leadingPartner_table_index, 0)
        # # Save the modified target document
        try:
            file_name_without_extension = os.path.splitext(source_path)[0]
            source_doc.save(f"{file_name_without_extension}_processed.docx")
            print("Wrong tables fixed correctly.")
            return True
        except Exception as e:
            print(f"Error saving the document: {e}")
            sys.exit(1)

# # Example usage
# if __name__ == "__main__":
#     if len(sys.argv) != 3:
#         print("Usage: python your_script_name.py <source_path> <target_path>")
#         sys.exit(1)
#     source_path = sys.argv[1]
#     target_path = sys.argv[2]
#     validate_document(source_path)
#     processed_path = source_path.replace(".docx", "_processed.docx")
#     copy_table_content_to_target(processed_path, target_path)
