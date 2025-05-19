from docx import Document
from docx.shared import RGBColor, Pt
from lxml import etree
import zipfile
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_UNDERLINE
from docx.enum.table import WD_ROW_HEIGHT
from docx.oxml.shared import OxmlElement, qn as oxml_qn
import io
import os
import sys
import re
from openai import OpenAI
from dotenv import load_dotenv


load_dotenv()

client = OpenAI(
    # This is the default and can be omitted
    api_key=os.getenv("OPENAI_API_KEY"),
)
template_Lawyer = """
        Steve Bobbins -> Steve Bobbins
        Steve Bobbins (Partner) -> Steve Bobbins (Partner)
        Steve Bobbins, Partner -> Steve Bobbins (Partner)
        Andrew Bobbins (partner) and Maria Prince (partner)  -> Andrew Bobbins (partner); Maria Prince (partner) 
        Steve Bobbins - Partner (Litigation) -> Steve Bobbins (Partner), Litigation
        John Jones, Ted Smith, Lisa Reilly -> John Jones; Ted Smith; Lisa Reilly
        John Jones (Partner, Litigation); Ted Smith, (Senior associate, Litigation); Lisa Reilly (Associate, Corporate) -> John Jones (Partner), Litigation; Ted Smith (Senior associate), Litigation; Lisa Reilly (Associate), Corporate
    """
template_advising = """
    '' -> ''
    George Jones of Lexington Legal Services– Counsel for Isabella Moore, Paul Baker of Bryant & Russell – Counsel for the Defendant, Horizon Legal Solutions – Advising the Defendant’s solicitors. -> George Jones of Lexington Legal Services - Counsel for Isabella Moore; Paul Baker of Bryant & Russell - Counsel for the Defendant; Horizon Legal Solutions - Advising the Defendant’s solicitors.
    George Jones of Lexington Legal Services– Civil action counsel for Frank Turner, Hugo Morris of Hartman & Fields – Copyright action counsel for Frank Turner, Penderby Jones – Defendant’s solicitors. -> George Jones of Lexington Legal Services - Civil action counsel for Frank Turner; Hugo Morris of Hartman & Fields - Copyright action counsel for Frank Turner; Penderby Jones - Defendant’s solicitors.
    Really Good  Law – Defendant’s solicitors -> Really Good  Law – Defendant’s solicitors
    Tony Blair of Sterling & Hayes, then Bobby Braithwaite of Sterling & Hayes, and Oliver Roberts of Aberdeen Chambers for trial – Counsel for Frank Lloyd, Jessica Jones, and Melanie Roberts, Good Law Solicitors – Advising the Defendant -> Tony Blair of Sterling & Hayes, then Bobby Braithwaite of Sterling & Hayes, and Oliver Roberts of Aberdeen Chambers for trial - Counsel for Frank Lloyd, Jessica Jones, and Melanie Roberts; Good Law Solicitors - Advising the Defendant
    Bobby Braithwaite of Lexington Legal Services– Counsel for Luminary Design Studio, Evergreen Legal Group – Advising the tenant -> Bobby Braithwaite of Lexington Legal Services - Counsel for Luminary Design Studio; Evergreen Legal Group - Advising the tenant
    Tony Blair of Lexington Legal Services– Counsel for Summit Environmental Solutions -> Tony Blair of Lexington Legal Services - Counsel for Summit Environmental Solutions
    Lawson & Associates – Advising the landlord -> Lawson & Associates - Advising the landlord
    Vanessa Smith (acting for the previous CRAYON fund), GMV (acting for the participants, represented by Hans van Meerten (prof CRAYON law) -> Vanessa Smith - acting for the previous CRAYON fund; GMV - acting for the participants, represented by Hans van Meerten (prof CRAYON law)
    B&B: our Belgian co-counsel; Stefan: counsel TWENTY FOUR; P&p: counsel for the management board members of TWENTY FOUR; -> B&B - our Belgian co-counsel; Stefan - counsel TWENTY FOUR; P&p - counsel for the management board members of TWENTY FOUR
    Allen & Overy Shearman for the insurance company\nPPP Rechtsbijstand representing the individuals -> Allen & Overy - Shearman for the insurance company; PPP - Rechtsbijstand representing the individuals    
    De Brauw (ANGELA SMITH and his team) advising SOMETHING.COM  -> De Brauw (ANGELA SMITH and his team) - advising SOMETHING.COM
    Smaller firm advising the employee\nFOURTY FIVE advising the insurance company of FOURTY SIX -> Smaller firm - advising the employee; FOURTY FIVE - advising the insurance company of FOURTY SIX
"""
# Define the template for formatting
def format_with_openai(input_str, template):
    # Construct the prompt with detailed instructions for formatting
    
    prompt = f"""
    In order to process strings, strings must be converted into a specific form according to their meaning.
    Given the following template rules:
    {template}
    
    Now, please format the following input according to these rules:
    {input_str}

    Return only the formatted output in the required format. No explanation is necessary.
    """
    
    # Request a response from the OpenAI API (using GPT-4)
    response = client.responses.create(
    model="gpt-3.5-turbo",
    instructions="You are a helpful assistant that can handle various strings.",
    input=prompt,
)
    if input_str is None or input_str == "":
        return ""
    # Return the formatted result from the response
    else:
        return response.output_text

practiceArea_text = ""
location_text = ""
result_path = ""

def delete_table_row(document: Document, table_index, row_index: int) -> None:
    """
    Deletes a specific row from a table in a Word document using python-docx's XML handling.
    
    Args:
        document (Document): The python-docx Document object
        table: The table object from which to delete the row
        row_index (int): Zero-based index of the row to delete
        
    Returns:
        None
        
    Raises:
        IndexError: If row_index is out of bounds
        ValueError: If the table has no rows or invalid input
    """
    table = document.tables[table_index]
    try:
        # Get the underlying XML table element
        table_element = table._tbl
        
        # Find all rows in the table using python-docx's XML methods
        rows = table_element.xpath('.//w:tr')
        
        if not rows:
            raise ValueError("Table has no rows to delete")
            
        if row_index < 0 or row_index >= len(rows):
            raise IndexError(f"Row index {row_index} is out of range. Table has {len(rows)} rows.")
        
        # Remove the specified row
        row_to_delete = rows[row_index]
        table_element.remove(row_to_delete)
        
    except Exception as e:
        raise RuntimeError(f"Failed to delete row: {str(e)}") from e
    
def write_to_specific_textbox(docx_file, textbox_index, new_text):
    """Writes a string to a specific textbox in a .docx file, updating both DrawingML and VML."""
    try:
        with zipfile.ZipFile(docx_file, 'r') as z:
            xml_content = z.read('word/document.xml')
            file_list = z.infolist()
            file_data = {info.filename: z.read(info.filename) for info in file_list}

        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
              'v': 'urn:schemas-microsoft-com:vml',
              'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
              'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
              'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}
        tree = etree.fromstring(xml_content, parser=etree.XMLParser(remove_blank_text=True))

        # Find both DrawingML and VML textboxes
        drawing_textboxes = tree.xpath('//wp:anchor/a:graphic/a:graphicData/wps:wsp/wps:txbx/w:txbxContent', namespaces=ns)
        vml_textboxes = tree.xpath('//v:shape/v:textbox/w:txbxContent', namespaces=ns)
        all_textboxes = drawing_textboxes + vml_textboxes

        if textbox_index < 0 or textbox_index >= len(all_textboxes):
            return False  # Textbox index out of range

        target_drawing_textbox = None
        target_vml_textbox = None

        if textbox_index < len(drawing_textboxes):
            target_drawing_textbox = drawing_textboxes[textbox_index]

        if textbox_index < len(vml_textboxes):
            target_vml_textbox = vml_textboxes[textbox_index]

        if target_drawing_textbox:
            text_elements_drawing = target_drawing_textbox.xpath('.//w:t', namespaces=ns)
            if text_elements_drawing:
                for element in text_elements_drawing:
                    element.text = ""
                text_elements_drawing[0].text = new_text
            else:
                p_element = target_drawing_textbox.xpath('.//w:p', namespaces=ns)[0]
                new_t = etree.SubElement(p_element, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
                new_t.text = new_text

        if target_vml_textbox:
            text_elements_vml = target_vml_textbox.xpath('.//w:t', namespaces=ns)
            if text_elements_vml:
                for element in text_elements_vml:
                    element.text = ""
                text_elements_vml[0].text = new_text
            else:
                p_element = target_vml_textbox.xpath('.//w:p', namespaces=ns)[0]
                new_t = etree.SubElement(p_element, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
                new_t.text = new_text

        xml_string = etree.tostring(tree, encoding='utf-8', xml_declaration=True)

        file_data['word/document.xml'] = xml_string

        with zipfile.ZipFile(docx_file, 'w', zipfile.ZIP_DEFLATED) as z:
            for filename, data in file_data.items():
                z.writestr(filename, data)

        return True

    except Exception as e:
        print(f"Error: {e}")
        return False
        
def clear_cell_content(cell):
    """
    Clears all paragraphs from a given cell.
    
    :param cell: The cell from which the content will be cleared
    """
    for paragraph in cell.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)


def extract_cell_text(source_doc, source_table_index, source_row_index, source_col_index):
    """
    Extracts text content from a specific table cell in a Word document using XML parsing.
    
    Args:
        source_doc: The Word document object (python-docx Document object)
        source_table_index: Index of the table in the document (0-based)
        source_row_index: Index of the row in the table (0-based)
        source_col_index: Index of the column in the row (0-based)
    
    Returns:
        A string containing the cell's text content, or None if extraction fails.
    """
    try:
        # ===== SOURCE CONTENT EXTRACTION =====
        try:
            source_xml_content = source_doc.part.blob
            source_tree = etree.fromstring(source_xml_content)
        except Exception as e:
            print(f"Error parsing source document XML: {str(e)}")
            return None

        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        
        # ===== SOURCE CELL VALIDATION =====
        source_tables = source_tree.xpath('//w:tbl', namespaces=ns)
        if source_table_index >= len(source_tables) or source_table_index < 0:
            print(f"Error: Source table index {source_table_index} out of range (0-{len(source_tables)-1})")
            return None

        source_rows = source_tables[source_table_index].xpath('.//w:tr', namespaces=ns)
        if source_row_index >= len(source_rows) or source_row_index < 0:
            print(f"Error: Source row index {source_row_index} out of range (0-{len(source_rows)-1})")
            return None

        source_cells = source_rows[source_row_index].xpath('.//w:tc', namespaces=ns)
        if source_col_index >= len(source_cells) or source_col_index < 0:
            print(f"Error: Source column index {source_col_index} out of range (0-{len(source_cells)-1})")
            return None

        # ===== TEXT EXTRACTION =====
        source_cell = source_cells[source_col_index]
        text_elements = source_cell.xpath('.//w:t', namespaces=ns)
        
        # Join all text elements with spaces
        cell_text = ' '.join([elem.text for elem in text_elements if elem.text])
        
        return cell_text.strip() if cell_text else None

    except Exception as e:
        print(f"Unexpected error during cell text extraction: {str(e)}")
        return None
    
def copy_cell_content_to_target_cell(source_doc, source_table_index, source_row_index, source_col_index,
                                    target_doc, target_table_index, target_row_index, target_col_index):
    """
    Copies cell content with formatting from one cell to another cell in a target Document object.
    """
    try:
        # ===== TARGET CELL VALIDATION =====
        if not hasattr(target_doc, 'tables') or target_table_index >= len(target_doc.tables) or target_table_index < 0:
            print(f"Error: Target table index {target_table_index} out of range (0-{len(target_doc.tables)-1})")
            return False
            
        target_table = target_doc.tables[target_table_index]
        if target_row_index >= len(target_table.rows) or target_row_index < 0:
            print(f"Error: Target row index {target_row_index} out of range (0-{len(target_table.rows)-1})")
            return False
            
        target_row = target_table.rows[target_row_index]
        if target_col_index >= len(target_row.cells) or target_col_index < 0:
            print(f"Error: Target column index {target_col_index} out of range (0-{len(target_row.cells)-1})")
            return False
            
        target_cell = target_row.cells[target_col_index]
        
        # ===== SOURCE CONTENT EXTRACTION =====
        try:
            source_xml_content = source_doc.part.blob
            source_tree = etree.fromstring(source_xml_content)
            
        except Exception as e:
            print(f"Error parsing source document XML: {str(e)}")
            return False

        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
              'wp': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
              'xml': 'http://www.w3.org/XML/1998/namespace'}
        
        # ===== SOURCE CELL VALIDATION =====
        source_tables = source_tree.xpath('//w:tbl', namespaces=ns)
        if source_table_index >= len(source_tables) or source_table_index < 0:
            print(f"Error: Source table index {source_table_index} out of range (0-{len(source_tables)-1})")
            return False

        source_rows = source_tables[source_table_index].xpath('.//w:tr', namespaces=ns)
        if source_row_index >= len(source_rows) or source_row_index < 0:
            print(f"Error: Source row index {source_row_index} out of range (0-{len(source_rows)-1})")
            return False

        source_cells = source_rows[source_row_index].xpath('.//w:tc', namespaces=ns)
        if source_col_index >= len(source_cells) or source_col_index < 0:
            print(f"Error: Source column index {source_col_index} out of range (0-{len(source_cells)-1})")
            return False

        source_cell = source_cells[source_col_index]
        source_paragraphs = source_cell.xpath('.//w:p', namespaces=ns)
        
        # ===== CONTENT COPYING =====
        target_cell.text = ""  # Clear target cell


        first_paragraph = True
        n = 0
        pre_left = ''
        pre_hanging = ''
        for i, paragraph in enumerate(source_paragraphs):
            
            # Create or get paragraph
            if first_paragraph:
                new_paragraph = target_cell.paragraphs[0]
                first_paragraph = False
            else:
                new_paragraph = target_cell.add_paragraph()
                    
            # Copy paragraph formatting
            ppr = paragraph.find('.//w:pPr', namespaces=ns)
            if ppr is not None:
                new_ppr = new_paragraph._element.get_or_add_pPr()
                
                numpr = ppr.find('.//w:numPr', namespaces=ns)
                if numpr is not None:
                    try:
                        ilvl = numpr.find('.//w:ilvl', namespaces=ns)
                        numId = numpr.find('.//w:numId', namespaces=ns)
                        if ilvl is not None and numId is not None:
                            # new_paragraph.paragraph_format.left_indent = Pt(18 * int(ilvl.get(qn('w:val'))) + 35)
                            new_numpr = etree.SubElement(new_ppr, qn('w:numPr'))
                            etree.SubElement(new_numpr, qn('w:ilvl'), {qn('w:val'): ilvl.get(qn('w:val'))})
                            etree.SubElement(new_numpr, qn('w:numId'), {qn('w:val'): '2'})
                    except Exception as e:
                        print(f"Warning: Could not copy numbering - {str(e)}")
                
                rpr = ppr.find('.//w:rPr', namespaces=ns)
                if rpr is not None:
                    color = rpr.find('.//w:color', namespaces=ns)
                    if color is not None:
                        if not new_paragraph.runs:
                            new_paragraph.add_run('')
                        
                        for run in new_paragraph.runs:
                            rPr = run._r.get_or_add_rPr()
                            etree.SubElement(rPr, qn('w:color'), {qn('w:val'): color.get(qn('w:val'))})
                
                # Copy the pStyle if it exists
                pstyle = ppr.find('.//w:pStyle', namespaces=ns)
                if pstyle is not None:
                    
                    new_style = etree.SubElement(new_ppr, qn('w:pStyle'))
                    style_name = pstyle.get(qn('w:val'))
                    new_style.set(qn('w:val'), style_name)
                    
                # indent handling
                ind = ppr.find('.//w:ind', namespaces=ns)
                if ind is not None and numpr is None:
                    new_ind = etree.SubElement(new_ppr, qn('w:ind'))
                    
                    left = ind.get(qn('w:left'))
                    # print(type(left))
                    if left is not None:
                        new_ind.set(qn('w:left'), left)
                        pre_left = left
                        n = i
                    hanging = ind.get(qn('w:hanging'))
                    if hanging is not None:
                        new_ind.set(qn('w:hanging'), hanging)
                        pre_hanging = hanging
                    
                
                if i == n + 1 and numpr is None:
                    # Create new indentation element
                    new_ind = etree.SubElement(new_ppr, qn('w:ind'))
                    if pre_left is not None:
                        new_ind.set(qn('w:left'), pre_left)
                    if pre_hanging is not None:
                        new_ind.set(qn('w:hanging'), pre_hanging)
                
                # Spacing handling
                spacing = ppr.find('.//w:spacing', namespaces=ns)
                is_last_paragraph = (i == len(source_paragraphs) - 1)
                is_empty = not paragraph.xpath('.//w:r/w:t', namespaces=ns)

                if spacing is not None:
                    # Space before
                    before = spacing.get(qn('w:before'))
                    if before is not None:
                        new_paragraph.paragraph_format.space_before = Pt(min(int(before)/20, 6))  # Max 6pt
                    
                    # Space after (none for last paragraph)
                    after = spacing.get(qn('w:after'))
                    if after is not None and not is_last_paragraph:
                        new_paragraph.paragraph_format.space_after = Pt(min(int(after)/20, 6))  # Max 6pt
                    else:
                        new_paragraph.paragraph_format.space_after = Pt(0)

                    # Line spacing
                    line = spacing.get(qn('w:line'))
                    if line is not None:
                        line_rule = spacing.get(qn('w:lineRule'), 'auto')
                        line_val = int(line)
                        if line_rule == 'exact':
                            new_paragraph.paragraph_format.line_spacing = Pt(line_val/20)
                        elif line_rule == 'atLeast':
                            new_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                            new_paragraph.paragraph_format.line_spacing = Pt(line_val/20)
                        else:  # auto or multiple
                            new_paragraph.paragraph_format.line_spacing = line_val/240  # 240 = 12pt*20
                else:
                    new_paragraph.paragraph_format.space_after = Pt(0)

            process_paragraph(paragraph, new_paragraph, ns)
            
        return True

    except Exception as e:
        print(f"Error: {e}")
        return False

def process_paragraph(paragraph, new_paragraph, ns):
    # Track both types of hyperlinks
    in_field_hyperlink = False
    in_modern_hyperlink = False
    hyperlink_data = {
        'url': None,
        'text_runs': [],
        'style': None
    }
    
    # Process all elements in original order
    for element in paragraph.xpath('./*', namespaces=ns):
        # Reset modern hyperlink flag at start of each element
        in_modern_hyperlink = False
        
        # Check for modern hyperlink (w:hyperlink element)
        if element.tag.endswith('}hyperlink'):
            process_modern_hyperlink(element, new_paragraph, ns)
            in_modern_hyperlink = True
            continue
        
        # Check for field-based hyperlink (legacy format)
        if element.tag.endswith('}r'):
            run = element
            fld_char_begin = run.find('.//w:fldChar[@w:fldCharType="begin"]', namespaces=ns)
            instr_text = run.find('.//w:instrText', namespaces=ns)
            fld_char_separate = run.find('.//w:fldChar[@w:fldCharType="separate"]', namespaces=ns)
            fld_char_end = run.find('.//w:fldChar[@w:fldCharType="end"]', namespaces=ns)
            text_element = run.find('.//w:t', namespaces=ns)
            
            if fld_char_begin is not None:
                in_field_hyperlink = True
                hyperlink_data = {'url': None, 'text_runs': [], 'style': None}
            
            elif in_field_hyperlink and instr_text is not None and 'HYPERLINK' in instr_text.text:
                match = re.search(r'HYPERLINK\s+"([^"]*)"', instr_text.text)
                if match:
                    hyperlink_data['url'] = match.group(1).strip()
            
            elif in_field_hyperlink and fld_char_separate is not None:
                continue
            
            elif in_field_hyperlink and text_element is not None:
                hyperlink_data['text_runs'].append(run)
            
            elif in_field_hyperlink and fld_char_end is not None:
                process_field_hyperlink(hyperlink_data, new_paragraph)
                in_field_hyperlink = False
                hyperlink_data = {'url': None, 'text_runs': [], 'style': None}
            
            # Process regular text (before, between, or after hyperlinks)
            if text_element is not None and text_element.text:
                # Only process if not part of any hyperlink
                if not in_field_hyperlink and not in_modern_hyperlink:
                    process_regular_run(run, new_paragraph, ns)

def process_modern_hyperlink(hyperlink, new_paragraph, ns):
    part = new_paragraph.part
    rels = part.rels
    
    # Generate unique rId
    next_rId = f"rId{len(rels) + 1}"
    
    # Add relationship
    rels.add_relationship(
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        "http://www.google.com",  # Default URL
        is_external=True,
        rId=next_rId
    )
    
    # Create hyperlink element
    hyperlink_copy = OxmlElement('w:hyperlink')
    hyperlink_copy.set(qn('r:id'), next_rId)
    
    # Copy all runs from original hyperlink
    for run in hyperlink.xpath('.//w:r', namespaces=ns):
        run_copy = OxmlElement('w:r')
        for child in run:
            run_copy.append(etree.fromstring(etree.tostring(child)))
        
        # Apply hyperlink formatting
        rPr = run_copy.find('w:rPr', namespaces=ns) or OxmlElement('w:rPr')
        
        # Set default font (Times New Roman, 11pt)
        font_name = OxmlElement('w:rFonts')
        font_name.set(qn('w:ascii'), 'Times New Roman')
        font_name.set(qn('w:hAnsi'), 'Times New Roman')
        rPr.insert(0, font_name)
        
        font_size = OxmlElement('w:sz')
        font_size.set(qn('w:val'), '22')  # 11pt = 22 half-points
        rPr.append(font_size)
        
        # Hyperlink style
        rStyle = OxmlElement('w:rStyle')
        rStyle.set(qn('w:val'), 'Hyperlink')
        rPr.insert(0, rStyle)
        
        # Blue text
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')
        rPr.append(color)
        
        # Black underline
        underline = OxmlElement('w:u')
        underline.set(qn('w:val'), 'single')
        underline.set(qn('w:color'), '000000')
        rPr.append(underline)
        
        # Preserve original formatting if exists
        original_rPr = run.find('w:rPr', namespaces=ns)
        if original_rPr is not None:
            # Copy bold/italic if present
            if original_rPr.find('.//w:b', namespaces=ns) is not None:
                bold = OxmlElement('w:b')
                rPr.append(bold)
            if original_rPr.find('.//w:i', namespaces=ns) is not None:
                italic = OxmlElement('w:i')
                rPr.append(italic)
        
        run_copy.insert(0, rPr)
        hyperlink_copy.append(run_copy)
    
    new_paragraph._p.append(hyperlink_copy)

def process_field_hyperlink(hyperlink_data, new_paragraph, ns):
    if not hyperlink_data['url'] or not hyperlink_data['text_runs']:
        return
    
    hyperlink_text = "".join([t.text for r in hyperlink_data['text_runs'] 
                            if (t := r.find('./w:t', namespaces=ns)) is not None])
    
    # Create begin field char
    begin_r = OxmlElement('w:r')
    begin_fldChar = OxmlElement('w:fldChar')
    begin_fldChar.set(qn('w:fldCharType'), 'begin')
    begin_r.append(begin_fldChar)
    new_paragraph._p.append(begin_r)
    
    # Create instrText run
    instr_r = OxmlElement('w:r')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = f' HYPERLINK "{hyperlink_data["url"]}" '
    instr_r.append(instr_text)
    new_paragraph._p.append(instr_r)
    
    # Create separate field char
    separate_r = OxmlElement('w:r')
    separate_fldChar = OxmlElement('w:fldChar')
    separate_fldChar.set(qn('w:fldCharType'), 'separate')
    separate_r.append(separate_fldChar)
    new_paragraph._p.append(separate_r)
    
    # Create hyperlink text run with style
    text_r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Add hyperlink style if exists
    if hyperlink_data['style']:
        rStyle = OxmlElement('w:rStyle')
        rStyle.set(qn('w:val'), hyperlink_data['style'])
        rPr.append(rStyle)
    
    # Add formatting elements
    font_size = OxmlElement('w:sz')
    font_size.set(qn('w:val'), '22')  # 11pt = 22 half-points
    rPr.append(font_size)
    
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')  # Blue
    rPr.append(color)
    
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)
    
    text_r.append(rPr)
    t = OxmlElement('w:t')
    t.text = hyperlink_text
    text_r.append(t)
    new_paragraph._p.append(text_r)
    
    # Create end field char with same formatting
    end_r = OxmlElement('w:r')
    end_rPr = OxmlElement('w:rPr')
    
    if hyperlink_data['style']:
        end_rStyle = OxmlElement('w:rStyle')
        end_rStyle.set(qn('w:val'), hyperlink_data['style'])
        end_rPr.append(end_rStyle)
    
    # Recreate formatting elements for end run
    end_font_size = OxmlElement('w:sz')
    end_font_size.set(qn('w:val'), '22')
    end_rPr.append(end_font_size)
    
    end_color = OxmlElement('w:color')
    end_color.set(qn('w:val'), '0000FF')
    end_rPr.append(end_color)
    
    end_underline = OxmlElement('w:u')
    end_underline.set(qn('w:val'), 'single')
    end_rPr.append(end_underline)
    
    end_r.append(end_rPr)
    end_fldChar = OxmlElement('w:fldChar')
    end_fldChar.set(qn('w:fldCharType'), 'end')
    end_r.append(end_fldChar)
    new_paragraph._p.append(end_r)

def process_regular_run(run, new_paragraph, ns):
    text_element = run.find('.//w:t', namespaces=ns)
    if text_element is not None and text_element.text:
        new_run = new_paragraph.add_run(text_element.text)
        
        # Default formatting
        new_run.font.name = "Times New Roman"
        new_run.font.size = Pt(11)
        
        # Copy run formatting
        rpr = run.find('.//w:rPr', namespaces=ns)
        if rpr is not None:
            # Color
            color_element = rpr.find('.//w:color', namespaces=ns)
            if color_element is not None and color_element.get(qn('w:val')):
                color_hex = color_element.get(qn('w:val'))
                try:
                    new_run.font.color.rgb = RGBColor.from_string(color_hex)
                except ValueError:
                    pass

            # Bold/Italic
            if rpr.find('.//w:b', namespaces=ns) is not None:
                new_run.bold = True
            if rpr.find('.//w:i', namespaces=ns) is not None:
                new_run.italic = True
                
            # Underline
            underline_element = rpr.find('.//w:u', namespaces=ns)
            if underline_element is not None:
                underline_val = underline_element.get(qn('w:val'))
                underline_mapping = {
                    'single': WD_UNDERLINE.SINGLE,
                    'double': WD_UNDERLINE.DOUBLE,
                    'thick': WD_UNDERLINE.THICK,
                    'dotted': WD_UNDERLINE.DOTTED,
                    'dash': WD_UNDERLINE.DASH,
                    'dotDash': WD_UNDERLINE.DOT_DASH,
                    'dotDotDash': WD_UNDERLINE.DOT_DOT_DASH,
                    'wave': WD_UNDERLINE.WAVY,
                    'none': WD_UNDERLINE.NONE,
                }
                new_run.font.underline = underline_mapping.get(underline_val, WD_UNDERLINE.SINGLE)
                
                # Underline color
                underline_color = underline_element.get(qn('w:color'))
                if underline_color:
                    try:
                        new_run.font.underline_color.rgb = RGBColor.from_string(underline_color)
                    except ValueError:
                        pass

def copy_row_formatting(source_row, target_row):
    """
    Copies the background color, borders, and other formatting from the source row to the target row.
    
    :param source_row: The row from which formatting will be copied
    :param target_row: The row to which formatting will be applied
    """
    namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    
    # Iterate through the cells of both rows and copy the background color and borders
    for source_cell, target_cell in zip(source_row.cells, target_row.cells):
        # Find the shading (background color) in the source cell
        source_shading = source_cell._element.find('.//w:shd', namespaces)
        
        if source_shading is not None:
            # Manually add the shading (background color) to the target cell
            target_cell_properties = target_cell._element.get_or_add_tcPr()
            target_shading = target_cell_properties.find('.//w:shd', namespaces)
            if target_shading is None:
                # Add shading if it doesn't already exist
                target_shading = etree.SubElement(target_cell_properties, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
            
            # Set the background color (fill) in the target cell
            target_shading.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', 
                               source_shading.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill'))

        # Copy borders
        source_borders = source_cell._element.find('.//w:tcBorders', namespaces)
        if source_borders is not None:
            target_cell_properties = target_cell._element.get_or_add_tcPr()
            target_borders = target_cell_properties.find('.//w:tcBorders', namespaces)
            if target_borders is None:
                target_borders = etree.SubElement(target_cell_properties, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders')

            # Copy each border element
            for border_element in source_borders.getchildren():
                tag = border_element.tag
                existing_target_border = target_borders.find(tag)
                if existing_target_border is not None:
                    target_borders.remove(existing_target_border)
                target_borders.append(etree.fromstring(etree.tostring(border_element)))

def insert_row_with_above_formatting_direct_xml(document: Document, table, row_index: int):
    """
    Inserts a new row with exact formatting from the row above by directly manipulating the document XML.
    """
    try:
        if row_index > 0:
            table_element = table._tbl
            rows = table_element.findall('.//w:tr', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})

            # Adjust row_index to be 0-based for finding the source row
            source_index = row_index

            if source_index >= len(rows):
                raise IndexError("Row index out of range.")

            source_row_xml = rows[source_index]
            new_row_xml = etree.fromstring(etree.tostring(source_row_xml))

            # Remove cell content (paragraphs) from the new row
            for tc in new_row_xml.findall('.//w:tc', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                for p in tc.findall('.//w:p', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                    tc.remove(p)

            # Determine the insertion position: one index after the source row
            insert_position = source_index + 3
            table_element.insert(insert_position, new_row_xml)

            return None
        else:
            print("Error: Cannot copy format from above row if row_index is 0.")
            return None
    except IndexError:
        print("Error: Row index out of range.")
        return None
    except Exception as e:
        print(f"Error: {e}")
        return None


def get_cell_background_color(cell):
    """
    Extracts the background color of a cell.
    
    """
    cell_xml = cell._element
    
    shading = cell_xml.find('.//w:shd', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
    
    if shading is not None:
        fill_color = shading.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
        return fill_color
    return None

def set_specific_dropdown_pre_display_text(doc, dropdown_index, text):
    """Sets the pre-displayed text of a specific dropdown in a Document object."""
    try:
        tree = doc.element  # Access the document's XML tree
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

        dropdowns = tree.xpath('//w:sdt[.//w:dropDownList]', namespaces=ns)

        if dropdown_index < 0 or dropdown_index >= len(dropdowns):
            print("Error: Dropdown index out of range.")
            return

        target_dropdown = dropdowns[dropdown_index]
        sdt_content = target_dropdown.find('.//w:sdtContent', namespaces=ns)

        # Clear existing text content
        for t_element in sdt_content.findall('.//w:t', namespaces=ns):
            sdt_content.remove(t_element)

        # Add new text content
        new_text_element = etree.Element(qn('w:t'))
        new_text_element.text = text
        sdt_content.append(new_text_element)

        # Re-parse the document
        doc.part._element = doc.element

    except Exception as e:
        print(f"Error: {e}")
    
    
def find_tables_with_specific_string(doc, search_string):
    """
    Searches all tables in the document and returns a list of table indexes
    where ANY cell contains the search string.

    Args:
        doc: The document object (python-docx Document).
        search_string: The string to search for in any cell of the tables.

    Returns:
        List of table indexes (0-based) where the string is found.
    """
    matching_table_indexes = []

    for i, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                if search_string in cell.text.strip():
                    matching_table_indexes.append(i)
                    break  # No need to check other cells in this table
            else:
                continue  # Only triggers if inner loop didn't break
            break  # Exit row loop if a match was found

    return matching_table_indexes

def delete_table_with_paragraphs(doc, table_index, num_paragraphs_above=0, num_paragraphs_below=0):
    """
    Deletes a specific table and its associated paragraphs (above and below).

    Args:
        doc (docx.Document): The Document object.
        table_index (int): The index of the table to delete.
        num_paragraphs_above (int): The number of paragraphs to delete above the table.
        num_paragraphs_below (int): The number of paragraphs to delete below the table.
    """
    try:
        table = doc.tables[table_index]
        table_xml = table._element
        body = doc.element.body

        paragraphs_above_xml = []
        element = table_xml.getprevious()
        count = 0
        while element is not None and count < num_paragraphs_above:
            if element.tag.endswith('p'):
                paragraphs_above_xml.insert(0, element)
                count += 1
            element = element.getprevious()

        paragraphs_below_xml = []
        element = table_xml.getnext()
        count = 0
        while element is not None and count < num_paragraphs_below:
            if element.tag.endswith('p'):
                paragraphs_below_xml.append(element)
                count += 1
            element = element.getnext()

        elements_to_remove = paragraphs_above_xml + [table_xml] + paragraphs_below_xml
        for element in elements_to_remove:
            if element is not None:
                body.remove(element)

        doc.part._element = doc.element

    except IndexError:
        print("Error: Table index out of range.")
    except Exception as e:
        print(f"Error: {e}")
        
def copy_table_with_paragraphs(source_doc, source_table_index, target_doc, target_index, num_paragraphs_above=0, num_paragraphs_below=0):
    """
    Copies a specific table and its associated paragraphs (above and below) to a specific position in a target document.
    """
    try:
        source_table = source_doc.tables[source_table_index]
        table_xml = source_table._element 

        paragraphs_above_xml = []
        element = table_xml.getprevious()
        count = 0
        while element is not None and count < num_paragraphs_above:
            if element.tag.endswith('p'):
                paragraphs_above_xml.insert(0, element)
                count += 1
            element = element.getprevious()

        paragraphs_below_xml = []
        element = table_xml.getnext()
        count = 0
        while element is not None and count < num_paragraphs_below:
            if element.tag.endswith('p'):
                paragraphs_below_xml.append(element)
                count += 1
            element = element.getnext()

        new_table_xml = etree.fromstring(etree.tostring(table_xml))

        target_body = target_doc.element.body

        if target_index < len(target_doc.tables):  
            target_table_xml = target_doc.tables[target_index]._element
            insert_before = target_table_xml
        else: 
            insert_before = None

        if insert_before is not None:
            for p_xml in paragraphs_above_xml:
                target_body.insert(target_body.index(insert_before), p_xml)
            target_body.insert(target_body.index(insert_before), new_table_xml)
            for p_xml in paragraphs_below_xml:
                target_body.insert(target_body.index(insert_before), p_xml)
        else:
            for p_xml in paragraphs_above_xml:
                target_body.append(p_xml)
            target_body.append(new_table_xml)
            for p_xml in paragraphs_below_xml:
                target_body.append(p_xml)

        target_doc.part._element = target_doc.element

    except IndexError:
        print("Error: Table index out of range.")
    except Exception as e:
        print(f"Error: {e}")
        
def add_page_break_before_table(doc, table_index):
    """
    Adds a page break before a specific table in a Document object.

    Args:
        doc (docx.Document): The Document object.
        table_index (int): The index of the table.
    """
    try:
        table = doc.tables[table_index]
        table_element = table._element
        body = doc.element.body

        new_paragraph = etree.Element(qn('w:p'))
        new_run = etree.SubElement(new_paragraph, qn('w:r'))
        new_br = etree.SubElement(new_run, qn('w:br'), {qn('w:type'): 'page'})

        body.insert(body.index(table_element), new_paragraph)

        doc.part._element = doc.element

    except IndexError:
        print("Error: Table index out of range.")
    except Exception as e:
        print(f"Error: {e}")

def add_single_line_space_before_table(doc, table_index):
    """
    Adds a single line space (a new empty paragraph) before a specific table
    in a Document object.

    Args:
        doc (docx.Document): The Document object.
        table_index (int): The index of the table.
    """
    try:
        table = doc.tables[table_index]
        table_element = table._element
        body = doc.element.body

        new_paragraph = etree.Element(qn('w:p'))
        # An empty <w:p> element will create a line space

        body.insert(body.index(table_element), new_paragraph)

        doc.part._element = doc.element

    except IndexError:
        print("Error: Table index out of range.")
    except Exception as e:
        print(f"Error: {e}")


def write_text_to_cell(doc, table_index, row_index, cell_index, text, font_size, bold=True, alignment="left", font_color=None):
    """
    Writes text with specified font size, style, bold formatting, alignment, and color
    to a specific cell.

    Args:
        doc (docx.Document): The Document object.
        table_index (int): The index of the table.
        row_index (int): The index of the row.
        cell_index (int): The index of the cell.
        text (str): The text to write.
        font_size (int): The font size in points.
        bold (bool): Whether to make the text bold (default: True).
        alignment (str): Text alignment ("left", "center", "right") (default: "left").
        font_color (tuple or None): The font color as an RGB tuple (e.g., (255, 0, 0) for red).
                                     Defaults to None (no explicit color set).
    """
    try:
        table = doc.tables[table_index]
        cell = table.rows[row_index].cells[cell_index]
        paragraph = cell.paragraphs[0]
        run = paragraph.clear().add_run(text)

        font = run.font
        font.name = "Calibri (Body)"
        font.size = Pt(font_size)
        font.bold = bold

        if font_color:
            font.color.rgb = RGBColor(*font_color)

        if alignment == "left":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif alignment == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == "right":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            print("Warning: Invalid alignment specified. Defaulting to left.")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    except IndexError:
        print("Error: Table, row, or cell index out of range.")
    except Exception as e:
        print(f"Error: {e}")
        
def copy_publishable_matter_to_target(source_doc, target_doc, source_table_index, target_table_index):
    
    source_publishableMatter_table = source_doc.tables[source_table_index]
    target_publishableMatter_table = target_doc.tables[target_table_index]
    # print(target_table_index)
    # insert_row_with_above_formatting(target_publishableMatter_table, 22)
    copy_cell_content_to_target_cell(source_doc, source_table_index, 2, 0, target_doc, target_table_index, 2, 0)
    copy_cell_content_to_target_cell(source_doc, source_table_index, 4, 0, target_doc, target_table_index, 4, 0)
    copy_cell_content_to_target_cell(source_doc, source_table_index, 6, 0, target_doc, target_table_index, 5, 1)
    copy_cell_content_to_target_cell(source_doc, source_table_index, 8, 0, target_doc, target_table_index, 8, 0)
    i = 0
    n = len(target_publishableMatter_table.rows)
    is_test = source_publishableMatter_table.cell(1, 0).text.strip()
    leadPartner_text = ""
    otherKeymembers_text = ""
    otherFirmadvising_text = ""
    
    # if "Name of client – this will be publishable" in is_test:
    leadPartner_text = extract_cell_text(source_doc, source_table_index, 10, 0)
    otherKeymembers_text = extract_cell_text(source_doc, source_table_index, 12, 0)
    otherFirmadvising_text = extract_cell_text(source_doc, source_table_index, 14, 0)
    # else:
    #     leadPartner_text = source_publishableMatter_table.cell(5, 0).text.strip()
    #     otherKeymembers_text = source_publishableMatter_table.cell(6, 0).text.strip()
    #     otherFirmadvising_text = source_publishableMatter_table.cell(7, 0).text.strip()
    # print(leadPartner_text)
    if leadPartner_text != "":
        leadPartner_text = format_with_openai(leadPartner_text, template_Lawyer)
    # print("LeadPartner_Text: ", leadPartner_text)
    # print(leadPartner_text)
    leadPartner_text_list = [item for item in leadPartner_text.split(';') if item]
    leadPartner_text_list_length = len(leadPartner_text_list)
    
    if otherKeymembers_text != "":
        otherKeymembers_text = format_with_openai(otherKeymembers_text, template_Lawyer)
    # print("Other Key Members: ", otherKeymembers_text, "\n")
    otherKeymembers_text_list = [item for item in otherKeymembers_text.split(';') if item]
    # print(otherKeymembers_text_list)
    otherKeymembers_text_list_length = len(otherKeymembers_text_list)
    # print(otherKeymembers_text_list)
    
    if otherFirmadvising_text != "":
        otherFirmadvising_text = format_with_openai(otherFirmadvising_text, template_advising)
    otherFirmadvising_text_list = [item for item in otherFirmadvising_text.split(';') if item]
    
    otherFirmadvising_text_list_length = len(otherFirmadvising_text_list)
    # print(otherFirmadvising_text)
    while i < n:
        cell_text = target_publishableMatter_table.cell(i, 0).text.strip()    
        if "Lead partner(s)" in cell_text:
            if leadPartner_text_list_length >= 3:
                for j in range(leadPartner_text_list_length - 2):
                    insert_row_with_above_formatting_direct_xml(target_doc, target_publishableMatter_table, i + j + 4)
                    n = n + 1
            for j in range(leadPartner_text_list_length):
                temp_leadPartner_text = leadPartner_text_list[j]
                try:
                    temp_name_text, temp_practiceArea_text = temp_leadPartner_text.split(',', 1)
                except ValueError:
                    temp_name_text = temp_leadPartner_text 
                    temp_practiceArea_text = ""  
                
                try:
                    target_publishableMatter_table.cell(i + j + 2, 0).text = temp_name_text
                    target_publishableMatter_table.cell(i + j + 2, 4).text = temp_practiceArea_text
                except IndexError:
                    print(f"IndexError: Unable to set text for row {i + j + 2}. Skipping this entry.")
                    continue
        
        if "Other key team members" in cell_text:
            if otherKeymembers_text_list_length >= 3:
                for j in range(otherKeymembers_text_list_length - 2):
                    insert_row_with_above_formatting_direct_xml(target_doc, target_publishableMatter_table, i + j + 3)
                    n = n + 1
            for j in range(otherKeymembers_text_list_length):
                temp_keymembers_text = otherKeymembers_text_list[j]
                # print(temp_keymembers_text)
                try:
                    temp_name_text, temp_practiceArea_text = temp_keymembers_text.split(',', 1)
                
                except ValueError:
                    temp_name_text = temp_keymembers_text 
                    temp_practiceArea_text = ""  
                try:
                    target_publishableMatter_table.cell(i + j + 2, 0).text = temp_name_text
                    target_publishableMatter_table.cell(i + j + 2, 4).text = temp_practiceArea_text
                except IndexError:
                    print(f"IndexError: Unable to set text for row {i + j + 2}. Skipping this entry.")
                    continue
                
        if "Other firms advising on the matter and their role(s)" in cell_text:
            if otherFirmadvising_text_list_length >= 2:
                for j in range(otherFirmadvising_text_list_length - 1):
                    # print(i, "   ", j)
                    insert_row_with_above_formatting_direct_xml(target_doc, target_publishableMatter_table, i + j + 2)
                    n = n + 1
            for j in range(otherFirmadvising_text_list_length):
                temp_otherFirmadvising_text = otherFirmadvising_text_list[j]
                
                try:
                    temp_name_text, temp_advising_text = temp_otherFirmadvising_text.split('-', 1)
                except ValueError:
                    temp_name_text = temp_otherFirmadvising_text
                    temp_advising_text = ""  
                
                # print(temp_name_text, temp_advising_text)
                try:
                    target_publishableMatter_table.cell(i + j + 2, 0).text = temp_name_text
                    target_publishableMatter_table.cell(i + j + 2, 4).text = temp_advising_text
                except IndexError:
                    print(f"IndexError: Unable to set text for row {i + j + 2}. Skipping this entry.")
                    continue
        i = i + 1   
            
    copy_cell_content_to_target_cell(source_doc, source_table_index, 16, 0, target_doc, target_table_index, n - 1, 3)       
    

async def chamber_l500_convert(source_docx_path, target_docx_path):
    # Load the source and target documents
    source_doc = Document(source_docx_path)
    target_doc = Document(target_docx_path)

    # Extract the Firm Name
    copy_cell_content_to_target_cell(source_doc, 0, 1, 0,
                                      target_doc, 0, 0, 0)
    # Extract the Practice Area
    target_practiceArea_table = source_doc.tables[1]
    global practiceArea_text 
    practiceArea_text= target_practiceArea_table.cell(1, 0).text.strip()

    # Extract the Location
    target_location_table = source_doc.tables[2]
    global location_text 
    location_text= target_location_table.cell(1, 0).text.strip()
    # print(location_text,"  ",  practiceArea_text)
    # Extract the Contact Details
    source_contactDetail_table = source_doc.tables[3]
    target_contactDetail_table = target_doc.tables[1]
    for i in range(len(target_contactDetail_table.rows) - 1, 1, -1):
        temp_row = target_contactDetail_table.rows[i]
        target_contactDetail_table._element.remove(temp_row._element)
    for i in range(2, len(source_contactDetail_table.rows)):
        if source_contactDetail_table.cell(i, 0).text.strip() != "":   
            if i < len(source_contactDetail_table.rows) - 3:
                source_row = target_contactDetail_table.rows[1]
                new_row = target_contactDetail_table.add_row()
                copy_row_formatting(source_row, new_row)
            copy_cell_content_to_target_cell(source_doc, 3, i, 0, target_doc, 1, i - 1, 0)
            copy_cell_content_to_target_cell(source_doc, 3, i, 1, target_doc, 1, i - 1, 2)
            copy_cell_content_to_target_cell(source_doc, 3, i, 2, target_doc, 1, i - 1, 3)
            
    
    # Extract the Department Name    
    copy_cell_content_to_target_cell(source_doc, 4, 1, 0, target_doc, 2, 0, 0)
    
    # Extract the Department Information
    copy_cell_content_to_target_cell(source_doc, 5, 1, 0, target_doc, 4, 0, 1)
    copy_cell_content_to_target_cell(source_doc, 5, 3, 0, target_doc, 4, 0, 4)
    
    # Extract the Heads of Team(Department)
    source_headsOfteam_table = source_doc.tables[6]
    target_headsOfteam_table = target_doc.tables[3]
    for i in range(len(target_headsOfteam_table.rows) - 1, 1, -1):
        temp_row = target_headsOfteam_table.rows[i]
        target_headsOfteam_table._element.remove(temp_row._element)
    for i in range(2, len(source_headsOfteam_table.rows)):
        if source_headsOfteam_table.cell(i, 0).text.strip() != "":   
            if i < len(source_headsOfteam_table.rows) - 3:
                source_row = target_headsOfteam_table.rows[1]
                new_row = target_headsOfteam_table.add_row()
                copy_row_formatting(source_row, new_row)
            copy_cell_content_to_target_cell(source_doc, 6, i, 0, target_doc, 3, i - 1, 0)
        
    # Extract the factor of best department
    copy_cell_content_to_target_cell(source_doc, 9, 1, 0, target_doc, 5, 0, 0)
    
    # Extract feedback in this practice area
    copy_cell_content_to_target_cell(source_doc, 11, 1, 0, target_doc, 7, 0, 0)
    
    # Extract the publishable clients
    publishableClients_table_indexlist = find_tables_with_specific_string(source_doc, search_string="D0 – PUBLISHABLE CLIENTS")
    if len(publishableClients_table_indexlist) != 0:
        publishableClients_table_index = publishableClients_table_indexlist[0]
        source_publishableClients_table = source_doc.tables[publishableClients_table_index]
        target_publishableClients_table = target_doc.tables[8]
        for i in range(len(target_publishableClients_table.rows) - 1, 1, -1):
            temp_row = target_publishableClients_table.rows[i]
            target_publishableClients_table._element.remove(temp_row._element)
        for i in range(2, len(source_publishableClients_table.rows)):
            if source_publishableClients_table.cell(i, 0).text.strip() != "":
                source_row = target_publishableClients_table.rows[1]
                new_row = target_publishableClients_table.add_row()
                copy_row_formatting(source_row, new_row)
                copy_cell_content_to_target_cell(source_doc, 12, i, 1, target_doc, 8, i - 1, 0)
                copy_cell_content_to_target_cell(source_doc, 12, i, 2, target_doc, 8, i - 1, 1)
    
            
    # Extract Confidential clients
    confidentialClients_table_indexlist = find_tables_with_specific_string(source_doc, search_string="E0 – CONFIDENTIAL CLIENTS")
    if len(confidentialClients_table_indexlist) != 0:
        confidentialClients_table_index = confidentialClients_table_indexlist[0]
        source_confidentialClients_table = source_doc.tables[confidentialClients_table_index]
        target_confidentialClients_table = target_doc.tables[9]
        for i in range(len(target_confidentialClients_table.rows) - 1, 1, -1):
            temp_row = target_confidentialClients_table.rows[i]
            target_confidentialClients_table._element.remove(temp_row._element)
        for i in range(2, len(source_confidentialClients_table.rows)):
            if source_confidentialClients_table.cell(i, 1).text.strip() != "":
                source_row = target_confidentialClients_table.rows[1]
                new_row = target_confidentialClients_table.add_row()
                copy_row_formatting(source_row, new_row)
                copy_cell_content_to_target_cell(source_doc, confidentialClients_table_index, i, 1, target_doc, 9, i - 1, 0)
                copy_cell_content_to_target_cell(source_doc, confidentialClients_table_index, i, 2, target_doc, 9, i - 1, 1)
                
    if len(publishableClients_table_indexlist) == 0:
        delete_table_with_paragraphs(target_doc, 8)
    if len(confidentialClients_table_indexlist) == 0:
        delete_table_with_paragraphs(target_doc, 9)
        
    # Extract the information of Ranked and Unranked lawyers
    # Extract Leading Partner Information
    leadingPartner_num = 0
    nextGenerationPartner_num = 0
    leadingAssoaciate_num = 0
    source_ranked_unrankedLawyers_table = source_doc.tables[8]
    source_ranked_unrankedLawyers_table_length = len(source_ranked_unrankedLawyers_table.rows)
    leadingPartner_indices = find_tables_with_specific_string(target_doc, search_string="Partner: leading partner")
    delete_table_with_paragraphs(target_doc, leadingPartner_indices[1])
    delete_table_with_paragraphs(target_doc, leadingPartner_indices[1])
    for i in range(source_ranked_unrankedLawyers_table_length - 2):
        if source_ranked_unrankedLawyers_table.cell(i + 2, 2).text.strip() == "Y":
            leadingPartner_num += 1
        if source_ranked_unrankedLawyers_table.cell(i + 2, 2).text.strip() == "N":
            # copy_cell_content_to_target_cell(source_doc, 8, i + 2, 0, target_doc, leadingAssociate_indices[i - leadingPartner_num], 2, 0)
            # copy_cell_content_to_target_cell(source_doc, 8, i + 2, 1, target_doc, leadingAssociate_indices[i - leadingPartner_num], 4, 0)
            leadingAssoaciate_num += 1
    # print(leadingPartner_num)
    for i in range(leadingPartner_num - 1):
        copy_table_with_paragraphs(target_doc, leadingPartner_indices[0], target_doc, leadingPartner_indices[0] + i)
    for i in range(leadingPartner_num):
        add_single_line_space_before_table(target_doc, leadingPartner_indices[0] + i)
        templeadingPartner_index = leadingPartner_indices[0] + i
        temp_text = "Partner: leading partner " + str(i + 1)
        write_text_to_cell(target_doc, templeadingPartner_index, 0, 0, temp_text, 14, alignment="left", font_color=(255, 255, 255))
        copy_cell_content_to_target_cell(source_doc, 8, i + 2, 0, target_doc, leadingPartner_indices[0] + i, 2, 0)
        copy_cell_content_to_target_cell(source_doc, 8, i + 2, 3, target_doc, leadingPartner_indices[0] + i, 2, 2)
        copy_cell_content_to_target_cell(source_doc, 8, i + 2, 1, target_doc, leadingPartner_indices[0] + i, 4, 0)
    if leadingPartner_num == 0:
        delete_table_with_paragraphs(target_doc, leadingPartner_indices[0])
        
    # Extract Leading Associate Information
    leadingAssociate_indices = find_tables_with_specific_string(target_doc, search_string="Associate: leading associate")
    delete_table_with_paragraphs(target_doc, leadingAssociate_indices[1])
    for i in range(leadingAssoaciate_num - 1):
        copy_table_with_paragraphs(target_doc, leadingAssociate_indices[0], target_doc, leadingAssociate_indices[0] + i)
    for i in range(leadingAssoaciate_num):
        add_single_line_space_before_table(target_doc, leadingAssociate_indices[0] + i)
        templeadingAssociate_index = leadingAssociate_indices[0] + i
        temp_text = "Associate: leading associate " + str(i + 1)
        write_text_to_cell(target_doc, templeadingAssociate_index, 0, 0, temp_text, 14, alignment="left", font_color=(255, 255, 255))
        copy_cell_content_to_target_cell(source_doc, 8, i + leadingPartner_num + 2, 0, target_doc, leadingAssociate_indices[0] + i, 2, 0)
        copy_cell_content_to_target_cell(source_doc, 8, i + leadingPartner_num + 2, 3, target_doc, leadingAssociate_indices[0] + i, 2, 2)
        copy_cell_content_to_target_cell(source_doc, 8, i + leadingPartner_num + 2, 1, target_doc, leadingAssociate_indices[0] + i, 4, 0)
    if leadingAssoaciate_num == 0:
        delete_table_with_paragraphs(target_doc, leadingAssociate_indices[0])
        
    nextGenerationPartner_indices = find_tables_with_specific_string(target_doc, search_string="Partner: next generation partner")
    ranked_unrankedLawyers_indices = leadingPartner_indices + nextGenerationPartner_indices
    
    # Extract Hires/Departures of partners in last 12 months
    source_hiresDepartures_table = source_doc.tables[7]
    target_hiresDepartures_table_indices = find_tables_with_specific_string(target_doc, search_string="Name (English)")
    target_hiresDepartures_table_index = target_hiresDepartures_table_indices[0]
    target_hiresDepartures_table = target_doc.tables[target_hiresDepartures_table_index]
    for i in range(len(target_hiresDepartures_table.rows) - 1, 1, -1):
        row = target_hiresDepartures_table.rows[i]
        target_hiresDepartures_table._element.remove(row._element)
    for i in range(2, len(source_hiresDepartures_table.rows) - 1):
        if source_hiresDepartures_table.cell(i, 0).text.strip() != "": 
            source_row = target_hiresDepartures_table.rows[1]
            new_row = target_hiresDepartures_table.add_row()
            copy_row_formatting(source_row, new_row)
            copy_cell_content_to_target_cell(source_doc, 7, i, 0, target_doc, target_hiresDepartures_table_index, i - 1, 0)
            copy_cell_content_to_target_cell(source_doc, 7, i, 1, target_doc, target_hiresDepartures_table_index, i - 1, 2)
            copy_cell_content_to_target_cell(source_doc, 7, i, 2, target_doc, target_hiresDepartures_table_index, i - 1, 3)
        
    # Extract Publishable Matter
    source_publishable_matter_indices = set(find_tables_with_specific_string(source_doc, search_string="Publishable Matter"))
    temp_indices = set(find_tables_with_specific_string(source_doc, search_string="D1 Name of client"))
    source_publishable_matter_indices = list(source_publishable_matter_indices.intersection(temp_indices))
    source_publishable_matter_indices.sort()
    publishable_matter_list_length = len(source_publishable_matter_indices)
    print(source_publishable_matter_indices)
    
    
    source_non_publishable_matter_indices = set(find_tables_with_specific_string(source_doc, search_string="Confidential Matter"))
    temp_indices = set(find_tables_with_specific_string(source_doc, search_string="E1 Name of client"))
    source_non_publishable_matter_indices = list(source_non_publishable_matter_indices.intersection(temp_indices))
    source_non_publishable_matter_indices.sort()
    non_publishable_matter_list_length = len(source_non_publishable_matter_indices)
    print(source_non_publishable_matter_indices)
    # print(publishable_matter_list_length)
    target_publishable_matter_indices = set(find_tables_with_specific_string(target_doc, search_string="Publishable matter"))
    temp_indices = set(find_tables_with_specific_string(target_doc, search_string="Name of client"))
    target_publishable_matter_indices = list(target_publishable_matter_indices.intersection(temp_indices))
    target_publishable_matter_indices.sort()
    target_publishable_matter_index = target_publishable_matter_indices[0]
    # print(target_publishable_matter_indices)
    if publishable_matter_list_length != 0 and non_publishable_matter_list_length != 0:
        for i in range(publishable_matter_list_length - 1):
            copy_table_with_paragraphs(target_doc, target_publishable_matter_index, target_doc, target_publishable_matter_index + i + 1, num_paragraphs_above=0, num_paragraphs_below=0)
        for i in range(non_publishable_matter_list_length - 1):
            copy_table_with_paragraphs(target_doc, target_publishable_matter_index + publishable_matter_list_length, target_doc, target_publishable_matter_index + publishable_matter_list_length + i + 1, num_paragraphs_above=0, num_paragraphs_below=0)
        for i in range(non_publishable_matter_list_length + publishable_matter_list_length - 2):
            add_page_break_before_table(target_doc, target_publishable_matter_index + 2 + i)
        for i in range(publishable_matter_list_length):
            temp_text = "Publishable Matter " + str(i + 1)
            write_text_to_cell(target_doc, i + target_publishable_matter_index, 0, 0, temp_text, 14, alignment="left")
            # copy_publishable_matter_to_target(source_doc, target_doc, 13 + i, target_publishable_matter_index + i)
        for i in range(non_publishable_matter_list_length):
            temp_text = "Non-Publishable Matter " + str(i + 1)
            write_text_to_cell(target_doc, i + target_publishable_matter_index + publishable_matter_list_length, 0, 0, temp_text, 14, alignment="left")
            # copy_publishable_matter_to_target(source_doc, target_doc, 14 + i + publishable_matter_list_length, target_publishable_matter_index + publishable_matter_list_length + i) 
        for i in range(publishable_matter_list_length):
            copy_publishable_matter_to_target(source_doc, target_doc, 13 + i, target_publishable_matter_index + i)
        for i in range(non_publishable_matter_list_length):
            copy_publishable_matter_to_target(source_doc, target_doc, 14 + i + publishable_matter_list_length, target_publishable_matter_index + publishable_matter_list_length + i)
    
    
    elif non_publishable_matter_list_length == 0:
        delete_table_with_paragraphs(target_doc, target_publishable_matter_indices[0] + 1)
        for i in range(publishable_matter_list_length - 1):
            copy_table_with_paragraphs(target_doc, target_publishable_matter_index, target_doc, target_publishable_matter_index + i + 1, num_paragraphs_above=0, num_paragraphs_below=0)
        for i in range(publishable_matter_list_length - 2):
            add_page_break_before_table(target_doc, target_publishable_matter_index + 2 + i)
        for i in range(publishable_matter_list_length):
            temp_text = "Publishable Matter " + str(i + 1)
            write_text_to_cell(target_doc, i + target_publishable_matter_index, 0, 0, temp_text, 14, alignment="left")
        for i in range(publishable_matter_list_length):
            copy_publishable_matter_to_target(source_doc, target_doc, 13 + i, target_publishable_matter_index + i)
        
        
    else: 
        delete_table_with_paragraphs(target_doc, target_publishable_matter_index)
        for i in range(non_publishable_matter_list_length - 1):
            copy_table_with_paragraphs(target_doc, target_publishable_matter_index, target_doc, target_publishable_matter_index + i + 1, num_paragraphs_above=0, num_paragraphs_below=0)
        for i in range(non_publishable_matter_list_length - 2):
            add_page_break_before_table(target_doc, target_publishable_matter_index + 2 + i)
        for i in range(non_publishable_matter_list_length):
            temp_text = "Non-Publishable Matter " + str(i + 1)
            write_text_to_cell(target_doc, i + target_publishable_matter_index + publishable_matter_list_length, 0, 0, temp_text, 14, alignment="left")
            # copy_publishable_matter_to_target(source_doc, target_doc, 14 + i + publishable_matter_list_length, target_publishable_matter_index + publishable_matter_list_length + i)
        for i in range(non_publishable_matter_list_length):
            copy_publishable_matter_to_target(source_doc, target_doc, 13 + i, target_publishable_matter_index + i)
    
    
    # Save the modified target document
    global result_path
    file_name_without_extension = os.path.splitext(source_docx_path)[0]
    file_name = f"{file_name_without_extension}_result.docx"
    target_doc.save(file_name)
    print(f"Content copied successfully. Result saved to: {file_name}")
    return file_name
    
    # Verify the file was created
    if not os.path.exists(result_path):
        raise FileNotFoundError(f"Failed to create output file at {result_path}")

async def validate_document(source_path):
    
    source_doc = Document(source_path)
    publishable_matter_indices = set(find_tables_with_specific_string(source_doc, search_string="Publishable Matter"))
    temp_indices = set(find_tables_with_specific_string(source_doc, search_string="Name of client"))
    publishable_matter_indices = list(publishable_matter_indices.intersection(temp_indices))
    publishable_matter_indices.sort()
    publishable_matter_list_length = len(publishable_matter_indices)
    # print(publishable_matter_indices)
    
    non_publishable_matter_indices1 = set(find_tables_with_specific_string(source_doc, search_string="Confidential Matter"))
    non_publishable_matter_indices = list(non_publishable_matter_indices1.intersection(temp_indices))
    non_publishable_matter_indices.sort()

    matter_indices = publishable_matter_indices + non_publishable_matter_indices
    matter_indices.sort()
    # print(matter_indices)
    
    matter_indices_length = len(matter_indices)
    # print(matter_indices)
    for i in range(matter_indices_length):
        matter_table_index = matter_indices[i]
        matter_table = source_doc.tables[matter_table_index]
        matter_table_length = len(matter_table.rows)
        top = 0
        for j in range(matter_table_length):
            cell_text = matter_table.cell(j, 0).text.strip()
            if "Publishable Matter" in cell_text or "Confidential Matter" in cell_text:
                top = j
                break
        if top != 0:
            for k in range(top):
                delete_table_row(source_doc, matter_table_index, 0)
    
    publisahble_matter_table_first_cell_texts = []
    for i in range(publishable_matter_list_length):
        publishable_matter_table_index = matter_indices[i]
        publishable_matter_table = source_doc.tables[publishable_matter_table_index]  
        publisahble_matter_table_first_cell_texts.append(publishable_matter_table.cell(0, 0).text.strip())
        
    publisahble_matter_table_numbers = [int(re.search(r'\d+', s).group()) for s in publisahble_matter_table_first_cell_texts]
    # print(publisahble_matter_table_numbers)
    publisahble_matter_table_missing_numbers = []
    if len(publisahble_matter_table_numbers) > 0:
        publisahble_matter_table_min_num = 1
        publisahble_matter_table_max_num = publisahble_matter_table_numbers[-1]
    
        # Generate the full expected consecutive sequence
        publisahble_matter_table_full_sequence = set(range(publisahble_matter_table_min_num, publisahble_matter_table_max_num + 1))
        publisahble_matter_table_actual_sequence = set(publisahble_matter_table_numbers)
        
        publisahble_matter_table_missing_numbers = sorted(publisahble_matter_table_full_sequence - publisahble_matter_table_actual_sequence)
    
        if publisahble_matter_table_missing_numbers:
            # print(publisahble_matter_table_missing_numbers)
            publishable_str = ', '.join(f"Publishable Matter {num}" for num in publisahble_matter_table_missing_numbers)
            # Construct the final message
            message = f"{publishable_str} was written wrong. Please rewrite it."
            return False, message
    
    non_publisahble_matter_table_first_cell_texts = []
    for i in range(publishable_matter_list_length, matter_indices_length):
        non_publishable_matter_table_index = matter_indices[i]
        non_publishable_matter_table = source_doc.tables[non_publishable_matter_table_index]  
        non_publisahble_matter_table_first_cell_texts.append(non_publishable_matter_table.cell(0, 0).text.strip())
    # print(non_publisahble_matter_table_first_cell_texts)
    non_publisahble_matter_table_numbers = [int(re.search(r'\d+', s).group()) for s in non_publisahble_matter_table_first_cell_texts]
    # print(non_publisahble_matter_table_numbers)
    non_publisahble_matter_table_missing_numbers = []
    if len(non_publisahble_matter_table_numbers) > 0:
        non_publisahble_matter_table_min_num = 1
        non_publisahble_matter_table_max_num = non_publisahble_matter_table_numbers[-1]
    
        # Generate the full expected consecutive sequence
        non_publisahble_matter_table_full_sequence = set(range(non_publisahble_matter_table_min_num, non_publisahble_matter_table_max_num + 1))
        non_publisahble_matter_table_actual_sequence = set(non_publisahble_matter_table_numbers)
    
        non_publisahble_matter_table_missing_numbers = sorted(non_publisahble_matter_table_full_sequence - non_publisahble_matter_table_actual_sequence)
        # print(non_publisahble_matter_table_missing_numbers)
        if non_publisahble_matter_table_missing_numbers:
            # print(non_publisahble_matter_table_missing_numbers)
            non_publishable_str = ', '.join(f"Confidential Matter {num}" for num in non_publisahble_matter_table_missing_numbers)
            # Construct the final message
            message = f"{non_publishable_str} was written wrong. Please rewrite it."
            return False, message          
        # # Save the modified target document
    if not publisahble_matter_table_missing_numbers and not non_publisahble_matter_table_missing_numbers:
        try:
            file_name_without_extension = os.path.splitext(source_path)[0]
            source_doc.save(f"{file_name_without_extension}_processed.docx")
            return True, "The document is valid."
        except Exception as e:
            print(f"Error saving the document: {e}")
            # sys.exit(1)
            return False

# Example usage
# if __name__ == "__main__":
#     if len(sys.argv) != 3:
#         print("Usage: python your_script_name.py <source_path> <target_path>")
#         sys.exit(1)
#     source_path = sys.argv[1]
#     target_path = sys.argv[2]
#     validate_document(source_path)
#     processed_path = source_path.replace(".docx", "_processed.docx")
#     copy_table_content_to_target(processed_path, target_path)

# print(location_text)
# Write Country and Practice Area
# write_to_specific_textbox(result_path, 0, location_text)
# write_to_specific_textbox(result_path, 1, practiceArea_text)